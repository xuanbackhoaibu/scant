import hashlib
import base64
import html
from pathlib import Path
from typing import Any, Dict, List, Optional
import docx
from docx.table import Table
from docx.text.paragraph import Paragraph
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status
from sqlalchemy.ext.asyncio import AsyncSession
from app.core.config import settings
from app.core.database import get_db
from app.models.entities import User, Template, TemplateVersion
from app.repositories.template_repo import template_repo, template_version_repo
from app.schemas.template import TemplateCreate, TemplateResponse, TemplateVersionResponse
from app.api.deps import get_current_user, get_current_user_optional
from app.services.templates.docx_template_analyzer import template_analyzer
from app.services.templates.template_reverse_engineering_service import template_reverse_engineer
from app.services.documents.docx_parser import docx_parser

router = APIRouter(prefix="/templates", tags=["templates"])


def _emu_to_px(value: Optional[int]) -> Optional[int]:
    if not value:
        return None
    return max(1, round(value / 9525))


def _length_to_px(value: Any) -> Optional[int]:
    if not value:
        return None
    try:
        return max(0, round(value.pt * 96 / 72))
    except Exception:
        return None


def _paragraph_style(paragraph: Paragraph) -> str:
    align_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
    parts = []
    if paragraph.alignment is not None:
        parts.append(f"text-align:{align_map.get(int(paragraph.alignment), 'left')}")
    pf = paragraph.paragraph_format
    left = _length_to_px(pf.left_indent)
    first = _length_to_px(pf.first_line_indent)
    before = _length_to_px(pf.space_before)
    after = _length_to_px(pf.space_after)
    if left:
        parts.append(f"margin-left:{left}px")
    if first:
        parts.append(f"text-indent:{first}px")
    if before is not None:
        parts.append(f"margin-top:{before}px")
    if after is not None:
        parts.append(f"margin-bottom:{after}px")
    return ";".join(parts)


def _run_html(run: Any) -> str:
    styles = []
    if run.bold:
        styles.append("font-weight:700")
    if run.italic:
        styles.append("font-style:italic")
    if run.underline:
        styles.append("text-decoration:underline")
    if run.font.size:
        styles.append(f"font-size:{run.font.size.pt:.1f}pt")
    if run.font.name:
        styles.append(f"font-family:'{html.escape(run.font.name)}', 'Times New Roman', serif")
    if run.font.color and run.font.color.rgb:
        styles.append(f"color:#{run.font.color.rgb}")

    chunks: List[str] = []
    for child in run._element:
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "t":
            chunks.append(html.escape(child.text or ""))
        elif tag == "tab":
            chunks.append('<span class="docx-tab"></span>')
        elif tag == "br":
            break_type = child.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type")
            if break_type == "page":
                chunks.append('<span class="docx-page-break"></span>')
            else:
                chunks.append("<br>")
        elif tag == "drawing":
            blips = child.xpath(".//a:blip")
            if not blips:
                continue
            rid = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            if not rid:
                continue
            rel = run.part.rels.get(rid)
            if not rel:
                continue
            blob = rel.target_part.blob
            content_type = rel.target_part.content_type or "image/png"
            encoded = base64.b64encode(blob).decode("ascii")
            extents = child.xpath(".//wp:extent")
            width = height = None
            if extents:
                width = _emu_to_px(int(extents[0].get("cx", "0")))
                height = _emu_to_px(int(extents[0].get("cy", "0")))
            size_style = []
            if width:
                size_style.append(f"width:{width}px")
            if height:
                size_style.append(f"height:{height}px")
            chunks.append(
                f'<img src="data:{content_type};base64,{encoded}" style="{";".join(size_style)}" />'
            )

    text = "".join(chunks) or html.escape(run.text or "")
    if not styles:
        return text
    return f'<span style="{";".join(styles)}">{text}</span>'


def _paragraph_html(paragraph: Paragraph) -> str:
    style = _paragraph_style(paragraph)
    class_name = "docx-paragraph docx-block"
    if paragraph.style and "heading" in paragraph.style.name.lower():
        class_name += " docx-heading"
    content = "".join(_run_html(run) for run in paragraph.runs)
    if not content.strip():
        content = "<br>"
    return f'<p class="{class_name}" style="{style}">{content}</p>'


def _cell_html(cell: Any) -> str:
    blocks = []
    for paragraph in cell.paragraphs:
        blocks.append(_paragraph_html(paragraph))
    for table in cell.tables:
        blocks.append(_table_html(table))
    return "".join(blocks) or "&nbsp;"


def _table_html(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cells.append(f"<td>{_cell_html(cell)}</td>")
        rows.append(f"<tr>{''.join(cells)}</tr>")
    return f'<table class="docx-table docx-block"><tbody>{"".join(rows)}</tbody></table>'


def _iter_docx_blocks(document: docx.Document):
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def _docx_to_preview_html(file_path: str) -> str:
    document = docx.Document(file_path)
    body = []
    for block in _iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            body.append(_paragraph_html(block))
        elif isinstance(block, Table):
            body.append(_table_html(block))
    return f"""<!doctype html>
<html>
<head>
<meta charset="utf-8" />
<style>
  html, body {{ margin:0; padding:0; background:#f1f5f9; }}
  body {{ font-family:'Times New Roman', Times, serif; color:#111827; }}
  .page {{ position:relative; width:794px; min-height:1123px; height:1123px; margin:24px auto; background:white; padding:76px 76px 96px 96px; box-sizing:border-box; box-shadow:0 1px 3px rgba(15,23,42,.16); overflow:hidden; }}
  .page-content {{ min-height:951px; max-height:951px; overflow:hidden; }}
  .page-number {{ position:absolute; left:96px; right:76px; bottom:34px; text-align:center; color:#475569; font-size:11pt; }}
  .docx-paragraph {{ margin:0 0 8px 0; line-height:1.35; font-size:13pt; white-space:normal; }}
  .docx-heading {{ font-weight:700; }}
  .docx-heading:not(:first-child) {{ margin-top:10px; }}
  .docx-tab {{ display:inline-block; width:48px; }}
  .docx-page-break {{ display:none; }}
  img {{ display:inline-block; max-width:100%; object-fit:contain; vertical-align:middle; }}
  .docx-table {{ border-collapse:collapse; margin:6px 0 10px 0; width:auto; max-width:100%; }}
  .docx-table td {{ border:1px solid #9ca3af; padding:3px 7px; vertical-align:middle; min-width:28px; }}
  .docx-table p {{ margin:0; }}
  #docx-source {{ display:none; }}
</style>
</head>
<body>
  <main id="docx-source">{"".join(body)}</main>
  <main id="docx-pages"></main>
  <script>
    (function() {{
      const source = document.getElementById("docx-source");
      const pages = document.getElementById("docx-pages");
      const blocks = Array.from(source.children);

      function createPage() {{
        const page = document.createElement("section");
        page.className = "page";
        const content = document.createElement("div");
        content.className = "page-content";
        const pageNumber = document.createElement("div");
        pageNumber.className = "page-number";
        page.appendChild(content);
        page.appendChild(pageNumber);
        pages.appendChild(page);
        return content;
      }}

      let current = createPage();
      for (const block of blocks) {{
        const node = block.cloneNode(true);
        current.appendChild(node);
        const hasExplicitPageBreak = !!node.querySelector(".docx-page-break");
        if (current.scrollHeight > current.clientHeight && current.children.length > 1) {{
          current.removeChild(node);
          current = createPage();
          current.appendChild(node);
        }}
        if (hasExplicitPageBreak) {{
          current = createPage();
        }}
      }}

      const lastPage = pages.lastElementChild;
      if (lastPage && !lastPage.querySelector(".page-content")?.children.length) {{
        lastPage.remove();
      }}

      const renderedPages = Array.from(document.querySelectorAll(".page"));
      renderedPages.forEach((page, index) => {{
        page.querySelector(".page-number").textContent = "Trang " + (index + 1) + " / " + renderedPages.length;
      }});
    }})();
  </script>
</body>
</html>"""


@router.get("")
async def list_templates(
    scope: str = "public",  # my, workspace, public, all
    category: Optional[str] = None,
    search: Optional[str] = None,
    current_user: User = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    from app.services.templates.template_library_service import template_library_service
    user_id = current_user.id if current_user else None
    return await template_library_service.list_templates(
        db=db,
        current_user_id=user_id,
        scope=scope,
        category=category,
        search=search,
    )


@router.post("/{template_id}/duplicate")
async def duplicate_template(
    template_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.services.templates.template_library_service import template_library_service
    return await template_library_service.duplicate_template(
        db=db,
        template_id=template_id,
        user_id=current_user.id,
        user_name=current_user.name
    )


@router.post("/{template_id}/publish")
async def publish_template(
    template_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.services.templates.template_library_service import template_library_service
    return await template_library_service.toggle_publish(
        db=db,
        template_id=template_id,
        user_id=current_user.id,
        publish=True
    )


@router.post("/{template_id}/unpublish")
async def unpublish_template(
    template_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.services.templates.template_library_service import template_library_service
    return await template_library_service.toggle_publish(
        db=db,
        template_id=template_id,
        user_id=current_user.id,
        publish=False
    )


@router.post("/{template_id}/use")
async def use_template(
    template_id: str,
    current_user: User = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    from app.services.templates.template_library_service import template_library_service
    await template_library_service.record_usage(db, template_id)
    return {"status": "success", "template_id": template_id}



@router.post("/reverse-engineer")
async def reverse_engineer_template(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
) -> Dict[str, Any]:
    """Decomposes an uploaded DOCX or PDF template into dynamic schema, fields, and styles."""
    filename = file.filename or "template.docx"
    ext = Path(filename).suffix.lower()
    if ext not in [".docx", ".doc", ".pdf"]:
        raise HTTPException(status_code=400, detail="Only .docx and .pdf files are supported")

    contents = await file.read()
    file_hash = hashlib.sha256(contents).hexdigest()
    stored_filename = f"tpl_rev_{file_hash[:12]}_{filename}"
    file_path = settings.TEMPLATE_DIR / stored_filename

    with open(file_path, "wb") as f:
        f.write(contents)

    if ext in [".docx", ".doc"]:
        schema = await template_reverse_engineer.reverse_engineer_docx(str(file_path))
    else:
        schema = await template_reverse_engineer.reverse_engineer_pdf(str(file_path))

    schema["stored_file_path"] = str(file_path)
    schema["original_filename"] = filename
    return schema


@router.post("/preview-docx")
async def preview_docx_template(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
) -> Dict[str, Any]:
    """Returns a lightweight, human-readable preview of an uploaded DOCX template."""
    filename = file.filename or "template.docx"
    ext = Path(filename).suffix.lower()
    if ext != ".docx":
        raise HTTPException(status_code=400, detail="Only .docx template preview is supported")

    contents = await file.read()
    file_hash = hashlib.sha256(contents).hexdigest()
    stored_filename = f"tpl_preview_{file_hash[:12]}_{filename}"
    file_path = settings.TEMPLATE_DIR / stored_filename

    with open(file_path, "wb") as f:
        f.write(contents)

    parsed = docx_parser.extract_document(str(file_path))
    paragraphs = [p for p in parsed.get("paragraphs", []) if p.get("text")]
    headings = parsed.get("headings", [])

    return {
        "original_filename": filename,
        "word_count": parsed.get("word_count", 0),
        "tables_count": parsed.get("tables_count", 0),
        "sections": parsed.get("sections", []),
        "headings": headings,
        "paragraphs": paragraphs,
        "preview_text": parsed.get("full_text", ""),
        "html_document": _docx_to_preview_html(str(file_path)),
    }


@router.post("/upload-docx", response_model=TemplateResponse)
async def upload_docx_template(
    name: str = Form(...),
    category: str = Form("business"),
    organization: str = Form(None),
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    filename = file.filename or "template.docx"
    ext = Path(filename).suffix.lower()
    if ext not in [".docx", ".doc"]:
        raise HTTPException(status_code=400, detail="Only .docx template files are supported")

    contents = await file.read()
    file_hash = hashlib.sha256(contents).hexdigest()
    stored_filename = f"tpl_{file_hash[:12]}_{filename}"
    file_path = settings.TEMPLATE_DIR / stored_filename

    with open(file_path, "wb") as f:
        f.write(contents)

    # Reverse engineer schema
    schema = await template_reverse_engineer.reverse_engineer_docx(str(file_path))

    template = await template_repo.create(db, obj_in={
        "user_id": current_user.id,
        "name": name,
        "category": category,
        "description": f"Mẫu văn bản trích xuất từ file {filename}",
        "is_system": False,
        "is_public": False,
        "organization": organization,
        "schema_json": schema,
    })

    version = await template_version_repo.create(db, obj_in={
        "template_id": template.id,
        "version_number": 1,
        "styles_json": schema.get("styles", {}),
        "placeholders_json": {
            "explicit": schema.get("explicit_placeholders", []),
            "detected": schema.get("fields", []),
        },
        "schema_json": schema,
        "file_path": str(file_path),
    })

    return TemplateResponse(
        id=template.id,
        user_id=template.user_id,
        name=template.name,
        category=template.category,
        description=template.description,
        is_system=template.is_system,
        is_public=template.is_public,
        organization=template.organization,
        created_at=template.created_at,
        updated_at=template.updated_at,
        latest_version=TemplateVersionResponse.model_validate(version)
    )
