import os
import math
import re
import struct
import unicodedata
import zlib
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional
import docx
from docx.shared import Inches, Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from app.core.config import settings
from app.services.citations.citation_formatter import citation_formatter


class DocxExporter:
    """High-fidelity DOCX Exporter generating editable academic documents conforming to university standards."""

    @classmethod
    def generate_docx(
        cls,
        report_title: str,
        topic_details: Dict[str, Any],
        sections: List[Any],
        sources: List[Any],
        document_settings: Optional[Dict[str, Any]] = None,
        include_cover: bool = True,
        include_toc: bool = True,
        include_references: bool = True,
        citation_style: str = "IEEE",
        template_path: Optional[str] = None,
        image_assets: Optional[Dict[str, Any]] = None,
    ) -> str:
        image_assets = image_assets or {}
        if template_path and Path(template_path).exists():
            doc = docx.Document(template_path)
            cls._fill_template_document(
                doc=doc,
                report_title=report_title,
                topic_details=topic_details,
                sections=sections,
                sources=sources,
                citation_style=citation_style,
                include_references=include_references,
                image_assets=image_assets,
            )
            filename = f"report_{os.urandom(6).hex()}.docx"
            out_path = settings.EXPORT_DIR / filename
            doc.save(str(out_path))
            return str(out_path)

        doc = docx.Document()

        # 1. Configure Standard A4 Page Setup & Academic Margins
        # Vietnam Academic Standard: Left 30mm, Right 20mm, Top 20mm, Bottom 20mm
        section = doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(30)
        section.right_margin = Mm(20)

        # Set Normal Style Font to Times New Roman 13pt
        style_normal = doc.styles["Normal"]
        style_normal.font.name = "Times New Roman"
        style_normal.font.size = Pt(13)
        style_normal.font.color.rgb = RGBColor(0, 0, 0)
        style_normal.paragraph_format.line_spacing = 1.5
        style_normal.paragraph_format.space_after = Pt(6)

        # 2. Cover Page
        if include_cover:
            cls._build_cover_page(doc, report_title, topic_details)
            doc.add_page_break()

        # 3. Table of Contents Header
        if include_toc:
            toc_heading = doc.add_paragraph()
            toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = toc_heading.add_run("MỤC LỤC")
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Times New Roman"

            # Add structured TOC items
            for sec in sections:
                p_toc = doc.add_paragraph()
                p_toc.paragraph_format.line_spacing = 1.2
                p_toc.paragraph_format.space_after = Pt(2)
                indent = (sec.level - 1) * 0.3
                p_toc.paragraph_format.left_indent = Inches(indent)

                r_title = p_toc.add_run(sec.title)
                r_title.font.name = "Times New Roman"
                r_title.font.size = Pt(12)
                if sec.level == 1:
                    r_title.bold = True

            doc.add_page_break()

        # 4. Report Sections Content
        has_written_section = False
        for sec in sections:
            # Heading formatting based on level
            if has_written_section and sec.level == 1 and cls._starts_new_page(sec.title):
                doc.add_page_break()

            p_head = doc.add_paragraph()
            p_head.paragraph_format.space_before = Pt(14)
            p_head.paragraph_format.space_after = Pt(8)
            p_head.paragraph_format.keep_with_next = True

            run_head = p_head.add_run(sec.title)
            run_head.font.name = "Times New Roman"

            if sec.level == 1:
                run_head.bold = True
                run_head.font.size = Pt(14)
                if "CHƯƠNG" in sec.title.upper():
                    p_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif sec.level == 2:
                run_head.bold = True
                run_head.font.size = Pt(13)
            else:
                run_head.bold = True
                run_head.italic = True
                run_head.font.size = Pt(13)

            # Paragraphs body
            insert_after = p_head
            body_text = sec.plain_text or ""
            # Strip the heading title if duplicated in plain_text
            lines = body_text.splitlines()
            if lines and lines[0].strip() == sec.title.strip():
                lines = lines[1:]

            cleaned_body = "\n".join(lines).strip()
            if cls._has_rich_image_nodes(getattr(sec, "content_json", None)):
                insert_after = cls._insert_tiptap_content_after(
                    insert_after,
                    getattr(sec, "content_json", None),
                    image_assets,
                    skip_title=sec.title,
                )
            elif cleaned_body:
                for trimmed in cls._split_content_blocks(cleaned_body):
                    insert_after = cls._insert_content_block_after(insert_after, trimmed)
            has_written_section = True

        # 5. References Section (Tài liệu tham khảo)
        if include_references and sources:
            doc.add_page_break()
            p_ref_head = doc.add_paragraph()
            p_ref_head.paragraph_format.space_before = Pt(14)
            p_ref_head.paragraph_format.space_after = Pt(8)
            r_ref_head = p_ref_head.add_run("TÀI LIỆU THAM KHẢO")
            r_ref_head.bold = True
            r_ref_head.font.size = Pt(14)
            r_ref_head.font.name = "Times New Roman"

            for idx, src in enumerate(sources, 1):
                src_dict = {
                    "title": src.title if hasattr(src, "title") else src.get("title"),
                    "authors": src.authors if hasattr(src, "authors") else src.get("authors"),
                    "publisher": src.publisher if hasattr(src, "publisher") else src.get("publisher"),
                    "published_date": src.published_date if hasattr(src, "published_date") else src.get("published_date"),
                    "url": src.url if hasattr(src, "url") else src.get("url"),
                }
                entry_text = citation_formatter.format_bibliography_entry(idx, src_dict, style=citation_style)
                
                p_entry = doc.add_paragraph()
                p_entry.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_entry.paragraph_format.line_spacing = 1.3
                p_entry.paragraph_format.space_after = Pt(4)
                p_entry.paragraph_format.left_indent = Inches(0.4)
                p_entry.paragraph_format.first_line_indent = Inches(-0.4)

                r_entry = p_entry.add_run(entry_text)
                r_entry.font.name = "Times New Roman"
                r_entry.font.size = Pt(12)

        # Save to disk
        filename = f"report_{os.urandom(6).hex()}.docx"
        out_path = settings.EXPORT_DIR / filename
        doc.save(str(out_path))

        return str(out_path)

    @classmethod
    def _fill_template_document(
        cls,
        doc: docx.Document,
        report_title: str,
        topic_details: Dict[str, Any],
        sections: List[Any],
        sources: List[Any],
        citation_style: str = "IEEE",
        include_references: bool = True,
        image_assets: Optional[Dict[str, Any]] = None,
    ) -> None:
        image_assets = image_assets or {}
        cls._replace_template_text(doc, report_title, topic_details)

        insert_after = cls._prepare_template_body_anchor(doc)
        insert_after = cls._insert_template_front_matter(insert_after, report_title, sections)

        for sec in sections:
            title = cls._clean_export_text(sec.title)
            if cls._is_intro_section(title):
                continue

            if cls._starts_new_page(title):
                cls._add_page_break_after(insert_after)

            heading_style = cls._heading_style_for_level(getattr(sec, "level", 1))
            insert_after = cls._insert_paragraph_after(insert_after, title, style=heading_style)
            insert_after.paragraph_format.keep_with_next = True
            insert_after.paragraph_format.space_before = Pt(0)
            insert_after.paragraph_format.space_after = Pt(6)
            cls._format_runs(insert_after, size=14 if sec.level == 1 else 13, bold=True)

            body_text = sec.plain_text or ""
            lines = body_text.splitlines()
            if lines and lines[0].strip() == sec.title.strip():
                lines = lines[1:]
            cleaned_blocks = [
                cls._clean_export_text(p_text)
                for p_text in cls._split_content_blocks("\n".join(lines))
            ]
            if cleaned_blocks and cleaned_blocks[0] == title:
                cleaned_blocks = cleaned_blocks[1:]
            if cls._has_rich_image_nodes(getattr(sec, "content_json", None)):
                insert_after = cls._insert_tiptap_content_after(
                    insert_after,
                    getattr(sec, "content_json", None),
                    image_assets,
                    skip_title=title,
                )
            elif cls._is_chapter_heading(title):
                insert_after = cls._insert_chapter_body_after(
                    insert_after,
                    title,
                    cleaned_blocks,
                    auto_subsections=not cls._chapter_has_child_sections(title, sections),
                )
            else:
                for trimmed in cleaned_blocks:
                    if not trimmed:
                        continue
                    insert_after = cls._insert_content_block_after(insert_after, trimmed)

        if include_references and sources:
            insert_after = cls._insert_paragraph_after(insert_after, "DANH MỤC TÀI LIỆU THAM KHẢO", style=None)
            insert_after.runs[0].bold = True
            insert_after.runs[0].font.name = "Times New Roman"
            insert_after.runs[0].font.size = Pt(14)
            for idx, src in enumerate(sources, 1):
                src_dict = {
                    "title": src.title if hasattr(src, "title") else src.get("title"),
                    "authors": src.authors if hasattr(src, "authors") else src.get("authors"),
                    "publisher": src.publisher if hasattr(src, "publisher") else src.get("publisher"),
                    "published_date": src.published_date if hasattr(src, "published_date") else src.get("published_date"),
                    "url": src.url if hasattr(src, "url") else src.get("url"),
                }
                entry_text = citation_formatter.format_bibliography_entry(idx, src_dict, style=citation_style)
                insert_after = cls._insert_paragraph_after(insert_after, entry_text, style=None)
                insert_after.paragraph_format.line_spacing = 1.3
                insert_after.paragraph_format.left_indent = Inches(0.4)
                insert_after.paragraph_format.first_line_indent = Inches(-0.4)
                for run in insert_after.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)

    @classmethod
    def _replace_template_text(cls, doc: docx.Document, report_title: str, topic: Dict[str, Any]) -> None:
        normalized_title = cls._normalize_report_title(report_title)
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text.upper().startswith("ĐỀ TÀI:"):
                cls._replace_paragraph_text_preserving_first_run(paragraph, f"ĐỀ TÀI: {normalized_title}")
                continue
            if text.upper() in {"(DATA CENTER)", "DATA CENTER"}:
                cls._replace_paragraph_text_preserving_first_run(paragraph, "")
                continue
            if text.startswith("Sinh viên thực hiện:") and (topic.get("student_name") or topic.get("lead_author")):
                cls._replace_paragraph_text_preserving_first_run(
                    paragraph,
                    f"Sinh viên thực hiện: {topic.get('student_name') or topic.get('lead_author') or ''}",
                )
            if text.startswith("Giáo viên hướng dẫn:") and topic.get("instructor"):
                cls._replace_paragraph_text_preserving_first_run(paragraph, f"Giáo viên hướng dẫn: {topic.get('instructor')}")

    @classmethod
    def _replace_paragraph_text_preserving_first_run(cls, paragraph, new_text: str) -> None:
        if not paragraph.runs:
            paragraph.add_run(new_text)
            return
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""

    @classmethod
    def _normalize_report_title(cls, report_title: str) -> str:
        title = cls._clean_export_text(report_title or "").strip()
        upper_title = title.upper()
        for prefix in ("ĐỀ TÀI:", "DE TAI:", "TOPIC:"):
            if upper_title.startswith(prefix):
                title = title[len(prefix):].strip()
                break
        for marker in ("Nếu bạn", "Neu ban", "tôi gợi ý", "toi goi y", "Yêu cầu", "Yeu cau", "Số trang", "So trang"):
            idx = title.lower().find(marker.lower())
            if idx > 8:
                title = title[:idx].strip()
        return title.strip(" “”\"")

    @classmethod
    def _insert_template_front_matter(cls, insert_after, report_title: str, sections: List[Any]):
        title = cls._normalize_report_title(report_title)
        intro_lines = cls._intro_lines_from_sections(sections, title)
        content_sections = [sec for sec in sections if not cls._is_intro_section(getattr(sec, "title", ""))]
        chapter_sections = [sec for sec in content_sections if not cls._is_references_section(getattr(sec, "title", ""))]
        table_count = sum(cls._count_markdown_tables(getattr(sec, "plain_text", "") or "") for sec in content_sections)

        front_matter = [
            ("LỜI NÓI ĐẦU", intro_lines, True),
            ("MỤC LỤC", [cls._toc_line(sec) for sec in content_sections], True),
            ("MỤC LỤC HÌNH ẢNH", ["Hình ảnh minh họa được giữ theo file mẫu và cập nhật khi người dùng chèn thêm hình trong Studio."], True),
            (
                "MỤC LỤC BẢNG BIỂU",
                [f"Bảng {idx}: Bảng phân tích và so sánh trong nội dung báo cáo" for idx in range(1, max(table_count, 1) + 1)],
                True,
            ),
            (
                "BẢNG CÁC TỪ VIẾT TẮT",
                ["AI: Trí tuệ nhân tạo", "CNTT: Công nghệ thông tin", "CPU: Bộ xử lý trung tâm", "RAM: Bộ nhớ truy cập ngẫu nhiên"],
                True,
            ),
        ]

        for heading, lines, page_break_after in front_matter:
            insert_after = cls._insert_paragraph_after(insert_after, heading, style="Heading 1")
            insert_after.paragraph_format.space_before = Pt(12)
            insert_after.paragraph_format.space_after = Pt(10)
            cls._format_runs(insert_after, size=14, bold=True)

            for line in lines:
                if not line:
                    continue
                insert_after = cls._insert_paragraph_after(insert_after, line, style=None)
                insert_after.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                insert_after.paragraph_format.line_spacing = 1.5
                insert_after.paragraph_format.first_line_indent = Inches(0.5) if heading == "LỜI NÓI ĐẦU" else None
                insert_after.paragraph_format.space_after = Pt(6)
                cls._format_runs(insert_after, size=13)

            if page_break_after:
                cls._add_page_break_after(insert_after)

        return insert_after

    @classmethod
    def _intro_lines_from_sections(cls, sections: List[Any], title: str) -> List[str]:
        intro_section = next((sec for sec in sections if cls._is_intro_section(getattr(sec, "title", ""))), None)
        if intro_section and getattr(intro_section, "plain_text", None):
            lines = intro_section.plain_text.splitlines()
            if lines and cls._is_intro_section(lines[0]):
                lines = lines[1:]
            paragraphs = [cls._clean_export_text(block) for block in cls._split_content_blocks("\n".join(lines))]
            paragraphs = [p for p in paragraphs if p]
            if paragraphs:
                return paragraphs[:6]

        return [
            f"Báo cáo trình bày nội dung nghiên cứu về đề tài \"{title}\" theo bố cục học thuật và yêu cầu của mẫu báo cáo Đại học Đại Nam.",
            "Nội dung được tổ chức theo ba phần chính: mở đầu, nội dung nghiên cứu theo chương và kết luận/tài liệu tham khảo.",
            "Trong quá trình biên soạn, báo cáo ưu tiên tính mạch lạc, căn cứ phân tích, bảng biểu so sánh và khả năng sử dụng trong học tập cũng như đánh giá môn học.",
        ]

    @classmethod
    def _toc_line(cls, sec: Any) -> str:
        title = cls._clean_export_text(getattr(sec, "title", "") or "")
        indent = "    " * max(0, (getattr(sec, "level", 1) or 1) - 1)
        return f"{indent}{title}\t..."

    @classmethod
    def _is_intro_section(cls, title: str) -> bool:
        return (title or "").strip().upper().startswith(("LỜI MỞ ĐẦU", "LỜI NÓI ĐẦU"))

    @classmethod
    def _is_references_section(cls, title: str) -> bool:
        return "TÀI LIỆU THAM KHẢO" in (title or "").strip().upper()

    @classmethod
    def _starts_new_page(cls, title: str) -> bool:
        text = (title or "").strip().upper()
        return text.startswith(("CHƯƠNG", "KẾT LUẬN", "TÀI LIỆU THAM KHẢO", "DANH MỤC TÀI LIỆU THAM KHẢO"))

    @classmethod
    def _is_chapter_heading(cls, title: str) -> bool:
        return (title or "").strip().upper().startswith("CHƯƠNG")

    @classmethod
    def _chapter_number(cls, title: str) -> Optional[int]:
        import re
        match = re.search(r"CHƯƠNG\s+(\d+)", (title or "").upper())
        return int(match.group(1)) if match else None

    @classmethod
    def _chapter_has_child_sections(cls, chapter_title: str, sections: List[Any]) -> bool:
        chapter_number = cls._chapter_number(chapter_title)
        if not chapter_number:
            return False
        prefix = f"{chapter_number}."
        for sec in sections:
            title = cls._clean_export_text(getattr(sec, "title", "") or "").strip()
            if title.startswith(prefix):
                return True
        return False

    @classmethod
    def _insert_chapter_body_after(cls, insert_after, chapter_title: str, blocks: List[str], auto_subsections: bool = True):
        clean_blocks = [block for block in blocks if block]
        if not clean_blocks:
            return insert_after

        if any(cls._looks_like_numbered_subheading(block) for block in clean_blocks):
            for block in clean_blocks:
                if cls._looks_like_numbered_subheading(block):
                    heading_level = cls._numbered_heading_level(block)
                    insert_after = cls._insert_paragraph_after(insert_after, block, style=cls._heading_style_for_level(heading_level))
                    insert_after.paragraph_format.keep_with_next = True
                    cls._format_runs(insert_after, size=13, bold=True)
                else:
                    insert_after = cls._insert_content_block_after(insert_after, block)
            return insert_after

        if not auto_subsections:
            for block in clean_blocks:
                insert_after = cls._insert_content_block_after(insert_after, block)
            return insert_after

        chapter_number = cls._chapter_number(chapter_title) or 1
        total_words = len(" ".join(clean_blocks).split())
        labels = cls._chapter_subsection_labels(chapter_number, chapter_title, total_words)
        groups = cls._split_blocks_into_groups(clean_blocks, len(labels))
        for idx, label in enumerate(labels, 1):
            insert_after = cls._insert_paragraph_after(insert_after, f"{chapter_number}.{idx}. {label}", style="Heading 2")
            insert_after.paragraph_format.keep_with_next = True
            insert_after.paragraph_format.space_before = Pt(8)
            insert_after.paragraph_format.space_after = Pt(4)
            cls._format_runs(insert_after, size=13, bold=True)

            subgroups = cls._split_blocks_into_groups(groups[idx - 1], 2)
            sublabels = cls._chapter_subsubsection_labels(idx)
            for sub_idx, subgroup in enumerate(subgroups, 1):
                if not subgroup:
                    continue
                if total_words >= 700:
                    insert_after = cls._insert_paragraph_after(
                        insert_after,
                        f"{chapter_number}.{idx}.{sub_idx}. {sublabels[sub_idx - 1]}",
                        style="Heading 3",
                    )
                    insert_after.paragraph_format.keep_with_next = True
                    insert_after.paragraph_format.space_before = Pt(6)
                    insert_after.paragraph_format.space_after = Pt(3)
                    cls._format_runs(insert_after, size=13, bold=True)
                for block in subgroup:
                    insert_after = cls._insert_content_block_after(insert_after, block)
        return insert_after

    @classmethod
    def _looks_like_numbered_subheading(cls, text: str) -> bool:
        import re
        stripped = (text or "").strip()
        return bool(re.match(r"^\d+\.\d+(\.\d+)?\s+", stripped)) and len(stripped.split()) <= 18

    @classmethod
    def _numbered_heading_level(cls, text: str) -> int:
        import re
        stripped = (text or "").strip()
        if re.match(r"^\d+\.\d+\.\d+\s+", stripped):
            return 3
        if re.match(r"^\d+\.\d+\s+", stripped):
            return 2
        return 1

    @classmethod
    def _split_blocks_into_groups(cls, blocks: List[str], group_count: int) -> List[List[str]]:
        groups: List[List[str]] = [[] for _ in range(group_count)]
        for idx, block in enumerate(blocks):
            groups[min(idx * group_count // max(1, len(blocks)), group_count - 1)].append(block)
        return groups

    @classmethod
    def _chapter_subsection_labels(cls, chapter_number: int, chapter_title: str, total_words: int) -> List[str]:
        specific = {
            1: ["Bối cảnh và mục tiêu nghiên cứu", "Phạm vi, đối tượng và phương pháp thực hiện", "Cấu trúc và định hướng triển khai báo cáo"],
            2: ["Cơ sở lý thuyết liên quan", "Các thành phần và công nghệ nền tảng", "Tiêu chí so sánh và đánh giá"],
            3: ["Phân tích yêu cầu và hiện trạng", "Thiết kế kiến trúc và mô hình xử lý", "Đánh giá phương án thiết kế"],
            4: ["Môi trường và quy trình triển khai", "Kết quả hiện thực hóa các thành phần chính", "Phân tích kết quả và bảng so sánh"],
            5: ["Mục tiêu và kịch bản kiểm thử", "Kết quả đánh giá hiệu năng", "Nhận xét, rủi ro và hướng cải thiện"],
            6: ["Tổng kết kết quả đạt được", "Hạn chế của báo cáo và hệ thống", "Hướng phát triển trong tương lai"],
        }
        labels = specific.get(chapter_number, ["Nội dung trọng tâm", "Phân tích chi tiết", "Đánh giá và kết luận"])
        return labels if total_words >= 900 else labels[:2]

    @classmethod
    def _chapter_subsubsection_labels(cls, subsection_index: int) -> List[str]:
        if subsection_index == 1:
            return ["Khái quát nội dung", "Ý nghĩa và phạm vi áp dụng"]
        if subsection_index == 2:
            return ["Phân tích chi tiết", "Nhận xét và đánh giá"]
        return ["Bảng biểu và minh chứng", "Kết luận tiểu mục"]

    @classmethod
    def _heading_style_for_level(cls, level: int) -> str:
        if level <= 1:
            return "Heading 1"
        if level == 2:
            return "Heading 2"
        return "Heading 3"

    @classmethod
    def _count_markdown_tables(cls, text: str) -> int:
        blocks = [block for block in cls._split_content_blocks(text or "") if "|" in block]
        return sum(1 for block in blocks if cls._parse_markdown_table(block))

    @classmethod
    def _split_content_blocks(cls, text: str) -> List[str]:
        blocks: List[str] = []
        paragraph_lines: List[str] = []
        table_lines: List[str] = []

        def flush_paragraph() -> None:
            nonlocal paragraph_lines
            block = "\n".join(paragraph_lines).strip()
            if block:
                blocks.append(block)
            paragraph_lines = []

        def flush_table() -> None:
            nonlocal table_lines
            if table_lines:
                blocks.append("\n".join(table_lines))
            table_lines = []

        for line in (text or "").splitlines():
            stripped = line.strip()
            is_table_row = cls._is_markdown_table_row(stripped)

            if is_table_row:
                flush_paragraph()
                table_lines.append(stripped)
                continue

            if not stripped:
                if table_lines:
                    continue
                flush_paragraph()
                continue

            flush_table()
            paragraph_lines.append(line)

        flush_table()
        flush_paragraph()
        return blocks

    @staticmethod
    def _is_markdown_table_row(text: str) -> bool:
        stripped = (text or "").strip()
        return stripped.startswith("|") and stripped.endswith("|") and "|" in stripped[1:-1]

    @classmethod
    def _has_rich_image_nodes(cls, content_json: Any) -> bool:
        found = False

        def visit(node: Any) -> None:
            nonlocal found
            if found or not isinstance(node, dict):
                return
            if node.get("type") == "image" and node.get("attrs", {}).get("assetId"):
                found = True
                return
            for child in node.get("content") or []:
                visit(child)

        visit(content_json)
        return found

    @classmethod
    def _node_text(cls, node: Dict[str, Any]) -> str:
        if node.get("type") == "text":
            return node.get("text") or ""
        return "".join(cls._node_text(child) for child in node.get("content") or [])

    @classmethod
    def _insert_tiptap_content_after(
        cls,
        insert_after,
        content_json: Any,
        image_assets: Dict[str, Any],
        skip_title: str = "",
    ):
        if not isinstance(content_json, dict):
            return insert_after
        for node in content_json.get("content") or []:
            if not isinstance(node, dict):
                continue
            node_type = node.get("type")
            attrs = node.get("attrs") or {}
            if node_type == "image":
                insert_after = cls._insert_asset_image_after(insert_after, attrs, image_assets)
                continue
            if node_type == "table":
                rows = cls._tiptap_table_rows(node)
                if rows:
                    table = cls._insert_table_after(insert_after, rows)
                    insert_after = cls._paragraph_after_table(table)
                continue
            text = cls._clean_export_text(cls._node_text(node))
            if not text or (skip_title and text.strip() == skip_title.strip()):
                continue
            if node_type == "heading":
                insert_after = cls._insert_paragraph_after(insert_after, text, style=cls._heading_style_for_level(int(attrs.get("level") or 2)))
                cls._format_runs(insert_after, size=14 if int(attrs.get("level") or 2) == 1 else 13, bold=True)
                continue
            insert_after = cls._insert_content_block_after(insert_after, text)
        return insert_after

    @classmethod
    def _tiptap_table_rows(cls, node: Dict[str, Any]) -> List[List[str]]:
        rows: List[List[str]] = []
        for row in node.get("content") or []:
            cells = []
            for cell in row.get("content") or []:
                cells.append(cls._clean_export_text(cls._node_text(cell)))
            if cells:
                rows.append(cells)
        return rows

    @classmethod
    def _insert_asset_image_after(cls, paragraph, attrs: Dict[str, Any], image_assets: Dict[str, Any]):
        asset_id = attrs.get("assetId")
        asset = image_assets.get(asset_id)
        storage_path = getattr(asset, "storage_path", None) if asset is not None else None
        if not storage_path or not Path(storage_path).exists():
            fallback = attrs.get("caption") or attrs.get("alt") or "Ảnh không còn tồn tại trong kho lưu trữ"
            return cls._insert_content_block_after(paragraph, f"[Ảnh: {fallback}]")

        picture_para = cls._insert_paragraph_after(paragraph, "", style=None)
        alignment = (attrs.get("alignment") or "center").lower()
        picture_para.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)

        width_px = attrs.get("width") or getattr(asset, "width", None) or 520
        width_inches = max(1.0, min(6.2, float(width_px) / 96.0))
        run = picture_para.add_run()
        run.add_picture(str(storage_path), width=Inches(width_inches))

        caption_text = cls._clean_export_text(attrs.get("caption") or "")
        if caption_text:
            caption = cls._insert_paragraph_after(picture_para, caption_text, style=None)
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.space_after = Pt(8)
            for run in caption.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.italic = True
            insert_after = caption
        else:
            insert_after = picture_para

        source = attrs.get("sourceName") or getattr(asset, "source_domain", None)
        license_value = attrs.get("license") or getattr(asset, "license", None)
        if source or license_value:
            source_text = "Nguồn ảnh: " + " · ".join([item for item in [source, license_value] if item])
            source_para = cls._insert_paragraph_after(insert_after, source_text, style=None)
            source_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in source_para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
                run.italic = True
            insert_after = source_para
        return insert_after

    @classmethod
    def _insert_content_block_after(cls, insert_after, text: str):
        visual_spec = cls._parse_visual_marker(text)
        if visual_spec:
            return cls._insert_visual_after(insert_after, visual_spec)

        table_rows = cls._parse_markdown_table(text)
        if table_rows:
            table = cls._insert_table_after(insert_after, table_rows)
            return cls._paragraph_after_table(table)

        insert_after = cls._insert_paragraph_after(insert_after, text, style=None)
        insert_after.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        insert_after.paragraph_format.line_spacing = 1.5
        insert_after.paragraph_format.first_line_indent = Inches(0.5)
        insert_after.paragraph_format.space_after = Pt(6)
        cls._format_runs(insert_after, size=13)
        return insert_after

    @classmethod
    def _parse_visual_marker(cls, text: str) -> Optional[Dict[str, Any]]:
        stripped = (text or "").strip()
        match = re.fullmatch(r"\[\[(CHART|IMAGE)\s*:(.*?)\]\]", stripped, flags=re.IGNORECASE | re.DOTALL)
        if not match:
            return None
        kind = match.group(1).lower()
        payload = match.group(2).strip()
        spec: Dict[str, Any] = {"kind": kind}
        for part in payload.split(";"):
            if "=" not in part:
                continue
            key, value = part.split("=", 1)
            spec[key.strip().lower()] = value.strip()
        if kind == "chart":
            labels = [item.strip() for item in spec.get("labels", "").split(",") if item.strip()]
            values = []
            for raw in spec.get("values", "").split(","):
                try:
                    values.append(float(raw.strip()))
                except Exception:
                    pass
            if not labels or not values:
                return None
            spec["labels"] = labels[:8]
            spec["values"] = values[: len(spec["labels"])]
            spec["chart_type"] = (spec.get("type") or "bar").lower()
        else:
            spec["title"] = spec.get("title") or "Ảnh minh họa"
            spec["prompt"] = spec.get("prompt") or spec["title"]
        return spec

    @classmethod
    def _insert_visual_after(cls, paragraph, spec: Dict[str, Any]):
        image_stream = (
            cls._build_chart_png(spec)
            if spec.get("kind") == "chart"
            else cls._build_illustration_png(spec)
        )
        title = cls._clean_export_text(spec.get("title") or spec.get("prompt") or "Minh họa")
        picture_para = cls._insert_paragraph_after(paragraph, "", style=None)
        picture_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = picture_para.add_run()
        run.add_picture(image_stream, width=Inches(5.8))

        caption_prefix = "Biểu đồ" if spec.get("kind") == "chart" else "Hình"
        caption = cls._insert_paragraph_after(picture_para, f"{caption_prefix}: {title}", style=None)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(8)
        for run in caption.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.italic = True
        return caption

    @classmethod
    def _build_chart_png(cls, spec: Dict[str, Any]) -> BytesIO:
        labels = spec.get("labels", [])
        values = spec.get("values", [])
        chart_type = spec.get("chart_type", "bar")
        title = spec.get("title") or "Biểu đồ thống kê"
        unit = spec.get("unit") or ""
        canvas = _PngCanvas(1100, 620, (255, 255, 255))
        canvas.rect(0, 0, 1099, 619, fill=(255, 255, 255), outline=(203, 213, 225))
        canvas.text(46, 32, title, scale=3, color=(15, 23, 42))
        canvas.text(48, 82, "Du lieu do AI tong hop tu noi dung bao cao", scale=2, color=(100, 116, 139))

        if chart_type in {"line", "area"}:
            cls._draw_line_chart(canvas, labels, values, unit)
        elif chart_type in {"pie", "donut"}:
            cls._draw_pie_chart(canvas, labels, values, unit)
        else:
            cls._draw_bar_chart(canvas, labels, values, unit)
        return canvas.to_png()

    @classmethod
    def _draw_bar_chart(cls, canvas, labels: List[str], values: List[float], unit: str) -> None:
        left, top, width, height = 90, 150, 900, 350
        max_value = max(values) if values else 1
        canvas.line(left, top, left, top + height, (71, 85, 105), 2)
        canvas.line(left, top + height, left + width, top + height, (71, 85, 105), 2)
        bar_gap = 18
        bar_w = max(34, int((width - bar_gap * (len(values) + 1)) / max(1, len(values))))
        colors = [(79, 70, 229), (14, 165, 233), (16, 185, 129), (245, 158, 11), (244, 63, 94), (100, 116, 139)]
        for idx, (label, value) in enumerate(zip(labels, values)):
            x = left + bar_gap + idx * (bar_w + bar_gap)
            bar_h = int((value / max_value) * (height - 42))
            y = top + height - bar_h
            canvas.rect(x, y, x + bar_w, top + height, fill=colors[idx % len(colors)])
            canvas.text(x, y - 24, cls._format_number(value, unit), scale=2, color=(15, 23, 42))
            canvas.text(x - 8, top + height + 18, label[:12], scale=2, color=(51, 65, 85))

    @classmethod
    def _draw_line_chart(cls, canvas, labels: List[str], values: List[float], unit: str) -> None:
        left, top, width, height = 90, 150, 900, 350
        max_value = max(values) if values else 1
        min_value = min(values) if values else 0
        span = max(1, max_value - min_value)
        canvas.line(left, top, left, top + height, (71, 85, 105), 2)
        canvas.line(left, top + height, left + width, top + height, (71, 85, 105), 2)
        points = []
        for idx, value in enumerate(values):
            x = left + int(idx * width / max(1, len(values) - 1))
            y = top + height - int(((value - min_value) / span) * (height - 42))
            points.append((x, y))
        for a, b in zip(points, points[1:]):
            canvas.line(a[0], a[1], b[0], b[1], (79, 70, 229), 4)
        for idx, (x, y) in enumerate(points):
            canvas.rect(x - 6, y - 6, x + 6, y + 6, fill=(14, 165, 233), outline=(15, 23, 42))
            canvas.text(x - 24, y - 30, cls._format_number(values[idx], unit), scale=2, color=(15, 23, 42))
            canvas.text(x - 24, top + height + 18, labels[idx][:12], scale=2, color=(51, 65, 85))

    @classmethod
    def _draw_pie_chart(cls, canvas, labels: List[str], values: List[float], unit: str) -> None:
        total = sum(values) or 1
        cx, cy, radius = 360, 335, 170
        colors = [(79, 70, 229), (14, 165, 233), (16, 185, 129), (245, 158, 11), (244, 63, 94), (100, 116, 139)]
        start = -90.0
        for idx, value in enumerate(values):
            end = start + 360.0 * value / total
            canvas.pie_slice(cx, cy, radius, start, end, colors[idx % len(colors)])
            start = end
        canvas.circle(cx, cy, radius, outline=(51, 65, 85))
        for idx, (label, value) in enumerate(zip(labels, values)):
            y = 180 + idx * 48
            color = colors[idx % len(colors)]
            canvas.rect(650, y, 684, y + 24, fill=color)
            canvas.text(700, y, f"{label[:18]}: {cls._format_number(value, unit)}", scale=2, color=(15, 23, 42))

    @classmethod
    def _build_illustration_png(cls, spec: Dict[str, Any]) -> BytesIO:
        title = spec.get("title") or "Ảnh minh họa"
        prompt = spec.get("prompt") or title
        canvas = _PngCanvas(1100, 620, (248, 250, 252))
        canvas.rect(0, 0, 1099, 619, fill=(248, 250, 252), outline=(203, 213, 225))
        canvas.text(46, 36, title, scale=3, color=(15, 23, 42))
        canvas.text(48, 90, "Anh minh hoa do AI tao theo yeu cau tai lieu", scale=2, color=(100, 116, 139))
        nodes = [
            (80, 210, 280, 340, "INPUT"),
            (345, 170, 575, 380, "AI ANALYSIS"),
            (640, 210, 840, 340, "REPORT"),
            (880, 210, 1040, 340, "DOCX"),
        ]
        colors = [(224, 231, 255), (219, 234, 254), (220, 252, 231), (254, 243, 199)]
        for idx, (x1, y1, x2, y2, label) in enumerate(nodes):
            canvas.rect(x1, y1, x2, y2, fill=colors[idx], outline=(79, 70, 229))
            canvas.text(x1 + 28, y1 + 50, label, scale=3 if idx != 1 else 2, color=(30, 41, 59))
            if idx < len(nodes) - 1:
                canvas.arrow(x2 + 12, (y1 + y2) // 2, nodes[idx + 1][0] - 12, (nodes[idx + 1][1] + nodes[idx + 1][3]) // 2, (79, 70, 229))
        canvas.text(92, 460, prompt[:82], scale=2, color=(51, 65, 85))
        return canvas.to_png()

    @classmethod
    def _format_number(cls, value: float, unit: str = "") -> str:
        rendered = f"{value:.1f}".rstrip("0").rstrip(".")
        return f"{rendered}{unit}"

    @classmethod
    def _parse_markdown_table(cls, text: str) -> Optional[List[List[str]]]:
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if len(lines) < 2 or not all("|" in line for line in lines):
            return None
        rows = []
        for line in lines:
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            if cells and all(cell.replace("-", "").replace(":", "").strip() == "" for cell in cells):
                continue
            rows.append(cells)
        if len(rows) < 2:
            return None
        column_count = len(rows[0])
        if column_count < 2 or any(len(row) != column_count for row in rows):
            return None
        return rows

    @classmethod
    def _insert_table_after(cls, paragraph, rows: List[List[str]]):
        usable_width = Inches(6.45)
        table = paragraph._parent.add_table(len(rows), len(rows[0]), usable_width)
        paragraph._p.addnext(table._tbl)
        table.autofit = False
        cls._set_table_layout(table, usable_width)
        preferred = cls._preferred_table_widths(len(rows[0]), usable_width)
        for r_idx, row_values in enumerate(rows):
            for c_idx, value in enumerate(row_values):
                cell = table.cell(r_idx, c_idx)
                cell.width = preferred[c_idx]
                cls._set_cell_width(cell, preferred[c_idx])
                cell.text = value
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(2)
                    cls._format_runs(p, size=12, bold=(r_idx == 0))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        return table

    @staticmethod
    def _preferred_table_widths(column_count: int, total_width) -> List[Any]:
        if column_count <= 1:
            return [total_width]
        if column_count == 2:
            ratios = [0.34, 0.66]
        elif column_count == 3:
            ratios = [0.24, 0.33, 0.43]
        elif column_count == 4:
            ratios = [0.20, 0.22, 0.22, 0.36]
        elif column_count == 5:
            ratios = [0.18, 0.20, 0.20, 0.20, 0.22]
        else:
            ratios = [1 / column_count] * column_count
        return [int(total_width * ratio) for ratio in ratios]

    @staticmethod
    def _set_table_layout(table, width) -> None:
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), str(DocxExporter._width_to_twips(width)))
        tbl_w.set(qn("w:type"), "dxa")

        layout = tbl_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")

    @staticmethod
    def _set_cell_width(cell, width) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:w"), str(DocxExporter._width_to_twips(width)))
        tc_w.set(qn("w:type"), "dxa")

    @staticmethod
    def _width_to_twips(width) -> int:
        if hasattr(width, "twips"):
            return int(width.twips)
        return int(int(width) / 635)

    @classmethod
    def _paragraph_after_table(cls, table):
        new_p = OxmlElement("w:p")
        table._tbl.addnext(new_p)
        return docx.text.paragraph.Paragraph(new_p, table._parent)

    @classmethod
    def _format_runs(cls, paragraph, size: int = 13, bold: bool = False) -> None:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            run.bold = bold

    @classmethod
    def _add_page_break_after(cls, paragraph) -> None:
        if not paragraph.runs:
            paragraph.add_run()
        paragraph.runs[-1].add_break(WD_BREAK.PAGE)

    @classmethod
    def _prepare_template_body_anchor(cls, doc: docx.Document):
        """
        Preserve the cover/scoring portion of an uploaded template and replace its
        placeholder body with generated content. This avoids appending AI content
        while stale sample sections such as "CHƯƠNG 1. TÊN CHƯƠNG" remain below it.
        """
        start_index = cls._find_template_body_start_index(doc)
        if start_index is None:
            insertion_index = cls._find_content_insertion_index(doc)
            return doc.paragraphs[insertion_index] if doc.paragraphs else doc.add_paragraph()

        anchor_index = max(0, start_index - 1)
        anchor = doc.paragraphs[anchor_index] if doc.paragraphs else doc.add_paragraph()
        body = doc.element.body
        start_element = doc.paragraphs[start_index]._p
        children = list(body)

        try:
            start_child_index = children.index(start_element)
        except ValueError:
            return anchor

        for child in children[start_child_index:]:
            if child.tag.endswith("}sectPr"):
                continue
            body.remove(child)

        return anchor

    @classmethod
    def _find_template_body_start_index(cls, doc: docx.Document) -> Optional[int]:
        markers = [
            "LỜI NÓI ĐẦU",
            "LỜI MỞ ĐẦU",
            "MỞ ĐẦU",
            "TÓM TẮT",
            "CHƯƠNG 1",
            "CHƯƠNG I",
            "I.",
        ]
        for idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip().upper()
            if not text:
                continue
            if any(text.startswith(marker) for marker in markers):
                return idx
        return None

    @classmethod
    def _clean_export_text(cls, text: str) -> str:
        cleaned_lines = []
        for raw_line in (text or "").splitlines():
            line = raw_line.strip()
            if not line or set(line) <= {"-", "—", "_"}:
                continue
            while line.startswith("#"):
                line = line[1:].strip()
            line = line.replace("**", "").replace("__", "")
            cleaned_lines.append(line)
        return "\n".join(cleaned_lines).strip()

    @classmethod
    def _find_content_insertion_index(cls, doc: docx.Document) -> int:
        markers = ["CHƯƠNG 1", "CHƯƠNG I", "LỜI NÓI ĐẦU", "MỤC LỤC"]
        for idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip().upper()
            if any(marker in text for marker in markers):
                return idx
        return max(0, len(doc.paragraphs) - 1)

    @classmethod
    def _insert_paragraph_after(cls, paragraph, text: str = "", style=None):
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = docx.text.paragraph.Paragraph(new_p, paragraph._parent)
        if style:
            try:
                new_para.style = style
            except Exception:
                pass
        if text:
            new_para.add_run(text)
        return new_para

    @classmethod
    def _build_cover_page(cls, doc: docx.Document, title: str, topic: Dict[str, Any]) -> None:
        university = topic.get("university") or "TRƯỜNG ĐẠI HỌC BÁCH KHOA HÀ NỘI"
        major = topic.get("major") or "VIỆN CÔNG NGHỆ THÔNG TIN VÀ TRUYỀN THÔNG"
        subject = topic.get("subject") or "BÁO CÁO BÀI TẬP LỚN / ĐỒ ÁN MÔN HỌC"
        student_name = topic.get("student_name") or "Nguyễn Văn A"
        student_id = topic.get("student_id") or "20210001"
        class_name = topic.get("class_name") or "K66-CNTT-01"
        instructor = topic.get("instructor") or "TS. Nguyễn Văn B"
        academic_year = topic.get("academic_year") or "Hà Nội, 2026"

        # University Header
        p_uni = doc.add_paragraph()
        p_uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_uni.paragraph_format.space_after = Pt(2)
        r_uni = p_uni.add_run(university.upper())
        r_uni.bold = True
        r_uni.font.size = Pt(13)
        r_uni.font.name = "Times New Roman"

        p_maj = doc.add_paragraph()
        p_maj.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_maj.paragraph_format.space_after = Pt(36)
        r_maj = p_maj.add_run(major.upper())
        r_maj.bold = True
        r_maj.font.size = Pt(12)
        r_maj.font.name = "Times New Roman"

        # Subject / Subtitle
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_after = Pt(18)
        r_sub = p_sub.add_run(subject.upper())
        r_sub.bold = True
        r_sub.font.size = Pt(14)
        r_sub.font.name = "Times New Roman"

        # Report Topic Title Box
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_after = Pt(72)
        r_title_label = p_title.add_run("ĐỀ TÀI:\n")
        r_title_label.bold = True
        r_title_label.font.size = Pt(13)
        r_title_label.font.name = "Times New Roman"

        r_title = p_title.add_run(title.upper())
        r_title.bold = True
        r_title.font.size = Pt(16)
        r_title.font.name = "Times New Roman"

        # Student & Instructor Table
        table = doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        info_rows = [
            ("Sinh viên thực hiện:", f"{student_name} - MSSV: {student_id}"),
            ("Lớp học phần:", class_name),
            ("Chuyên ngành:", major),
            ("Giảng viên hướng dẫn:", instructor),
        ]

        for i, (label, val) in enumerate(info_rows):
            row = table.rows[i]
            # Left cell (label)
            p_lbl = row.cells[0].paragraphs[0]
            p_lbl.paragraph_format.line_spacing = 1.3
            p_lbl.paragraph_format.space_after = Pt(4)
            r_l = p_lbl.add_run(label)
            r_l.font.name = "Times New Roman"
            r_l.font.size = Pt(12)
            r_l.bold = True

            # Right cell (value)
            p_v = row.cells[1].paragraphs[0]
            p_v.paragraph_format.line_spacing = 1.3
            p_v.paragraph_format.space_after = Pt(4)
            r_v = p_v.add_run(val)
            r_v.font.name = "Times New Roman"
            r_v.font.size = Pt(12)

        # Footer Date
        p_footer = doc.add_paragraph()
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_footer.paragraph_format.space_before = Pt(72)
        r_date = p_footer.add_run(academic_year)
        r_date.bold = True
        r_date.font.size = Pt(12)
        r_date.font.name = "Times New Roman"


docx_exporter = DocxExporter()


class _PngCanvas:
    _FONT = {
        "A": ["01110", "10001", "10001", "11111", "10001", "10001", "10001"],
        "B": ["11110", "10001", "10001", "11110", "10001", "10001", "11110"],
        "C": ["01111", "10000", "10000", "10000", "10000", "10000", "01111"],
        "D": ["11110", "10001", "10001", "10001", "10001", "10001", "11110"],
        "E": ["11111", "10000", "10000", "11110", "10000", "10000", "11111"],
        "F": ["11111", "10000", "10000", "11110", "10000", "10000", "10000"],
        "G": ["01111", "10000", "10000", "10011", "10001", "10001", "01111"],
        "H": ["10001", "10001", "10001", "11111", "10001", "10001", "10001"],
        "I": ["11111", "00100", "00100", "00100", "00100", "00100", "11111"],
        "J": ["00111", "00010", "00010", "00010", "10010", "10010", "01100"],
        "K": ["10001", "10010", "10100", "11000", "10100", "10010", "10001"],
        "L": ["10000", "10000", "10000", "10000", "10000", "10000", "11111"],
        "M": ["10001", "11011", "10101", "10101", "10001", "10001", "10001"],
        "N": ["10001", "11001", "10101", "10011", "10001", "10001", "10001"],
        "O": ["01110", "10001", "10001", "10001", "10001", "10001", "01110"],
        "P": ["11110", "10001", "10001", "11110", "10000", "10000", "10000"],
        "Q": ["01110", "10001", "10001", "10001", "10101", "10010", "01101"],
        "R": ["11110", "10001", "10001", "11110", "10100", "10010", "10001"],
        "S": ["01111", "10000", "10000", "01110", "00001", "00001", "11110"],
        "T": ["11111", "00100", "00100", "00100", "00100", "00100", "00100"],
        "U": ["10001", "10001", "10001", "10001", "10001", "10001", "01110"],
        "V": ["10001", "10001", "10001", "10001", "10001", "01010", "00100"],
        "W": ["10001", "10001", "10001", "10101", "10101", "10101", "01010"],
        "X": ["10001", "10001", "01010", "00100", "01010", "10001", "10001"],
        "Y": ["10001", "10001", "01010", "00100", "00100", "00100", "00100"],
        "Z": ["11111", "00001", "00010", "00100", "01000", "10000", "11111"],
        "0": ["01110", "10001", "10011", "10101", "11001", "10001", "01110"],
        "1": ["00100", "01100", "00100", "00100", "00100", "00100", "01110"],
        "2": ["01110", "10001", "00001", "00010", "00100", "01000", "11111"],
        "3": ["11110", "00001", "00001", "01110", "00001", "00001", "11110"],
        "4": ["00010", "00110", "01010", "10010", "11111", "00010", "00010"],
        "5": ["11111", "10000", "10000", "11110", "00001", "00001", "11110"],
        "6": ["01110", "10000", "10000", "11110", "10001", "10001", "01110"],
        "7": ["11111", "00001", "00010", "00100", "01000", "01000", "01000"],
        "8": ["01110", "10001", "10001", "01110", "10001", "10001", "01110"],
        "9": ["01110", "10001", "10001", "01111", "00001", "00001", "01110"],
        " ": ["00000", "00000", "00000", "00000", "00000", "00000", "00000"],
        ".": ["00000", "00000", "00000", "00000", "00000", "01100", "01100"],
        ",": ["00000", "00000", "00000", "00000", "01100", "00100", "01000"],
        ":": ["00000", "01100", "01100", "00000", "01100", "01100", "00000"],
        "-": ["00000", "00000", "00000", "11111", "00000", "00000", "00000"],
        "/": ["00001", "00010", "00010", "00100", "01000", "01000", "10000"],
        "%": ["11001", "11010", "00010", "00100", "01000", "01011", "10011"],
    }

    def __init__(self, width: int, height: int, background=(255, 255, 255)):
        self.width = width
        self.height = height
        self.pixels = bytearray(background * width * height)

    def _set(self, x: int, y: int, color) -> None:
        if 0 <= x < self.width and 0 <= y < self.height:
            idx = (y * self.width + x) * 3
            self.pixels[idx:idx + 3] = bytes(color)

    def rect(self, x1: int, y1: int, x2: int, y2: int, fill=None, outline=None) -> None:
        x1, x2 = sorted((max(0, x1), min(self.width - 1, x2)))
        y1, y2 = sorted((max(0, y1), min(self.height - 1, y2)))
        if fill:
            for y in range(y1, y2 + 1):
                start = (y * self.width + x1) * 3
                end = (y * self.width + x2 + 1) * 3
                self.pixels[start:end] = bytes(fill) * (x2 - x1 + 1)
        if outline:
            self.line(x1, y1, x2, y1, outline)
            self.line(x2, y1, x2, y2, outline)
            self.line(x2, y2, x1, y2, outline)
            self.line(x1, y2, x1, y1, outline)

    def line(self, x1: int, y1: int, x2: int, y2: int, color, thickness: int = 1) -> None:
        dx = abs(x2 - x1)
        dy = -abs(y2 - y1)
        sx = 1 if x1 < x2 else -1
        sy = 1 if y1 < y2 else -1
        err = dx + dy
        x, y = x1, y1
        while True:
            half = max(0, thickness // 2)
            for ox in range(-half, half + 1):
                for oy in range(-half, half + 1):
                    self._set(x + ox, y + oy, color)
            if x == x2 and y == y2:
                break
            e2 = 2 * err
            if e2 >= dy:
                err += dy
                x += sx
            if e2 <= dx:
                err += dx
                y += sy

    def arrow(self, x1: int, y1: int, x2: int, y2: int, color) -> None:
        self.line(x1, y1, x2, y2, color, 4)
        angle = math.atan2(y2 - y1, x2 - x1)
        for delta in (2.55, -2.55):
            ax = int(x2 - 20 * math.cos(angle + delta))
            ay = int(y2 - 20 * math.sin(angle + delta))
            self.line(x2, y2, ax, ay, color, 4)

    def circle(self, cx: int, cy: int, radius: int, outline=None, fill=None) -> None:
        r2 = radius * radius
        for y in range(cy - radius, cy + radius + 1):
            for x in range(cx - radius, cx + radius + 1):
                d2 = (x - cx) ** 2 + (y - cy) ** 2
                if fill and d2 <= r2:
                    self._set(x, y, fill)
                if outline and abs(d2 - r2) < radius * 2:
                    self._set(x, y, outline)

    def pie_slice(self, cx: int, cy: int, radius: int, start_deg: float, end_deg: float, color) -> None:
        while end_deg < start_deg:
            end_deg += 360
        for y in range(cy - radius, cy + radius + 1):
            for x in range(cx - radius, cx + radius + 1):
                dx = x - cx
                dy = y - cy
                if dx * dx + dy * dy > radius * radius:
                    continue
                angle = math.degrees(math.atan2(dy, dx))
                if angle < -90:
                    angle += 360
                if start_deg <= angle <= end_deg:
                    self._set(x, y, color)

    def text(self, x: int, y: int, value: str, scale: int = 2, color=(15, 23, 42)) -> None:
        safe = unicodedata.normalize("NFD", value or "").encode("ascii", "ignore").decode("ascii").upper()
        cursor = x
        for char in safe:
            glyph = self._FONT.get(char, self._FONT.get(" "))
            for row_idx, row in enumerate(glyph):
                for col_idx, bit in enumerate(row):
                    if bit == "1":
                        self.rect(
                            cursor + col_idx * scale,
                            y + row_idx * scale,
                            cursor + (col_idx + 1) * scale - 1,
                            y + (row_idx + 1) * scale - 1,
                            fill=color,
                        )
            cursor += 6 * scale
            if cursor > self.width - 24:
                break

    def to_png(self) -> BytesIO:
        raw = bytearray()
        stride = self.width * 3
        for y in range(self.height):
            raw.append(0)
            raw.extend(self.pixels[y * stride:(y + 1) * stride])
        compressed = zlib.compress(bytes(raw), 9)

        def chunk(tag: bytes, data: bytes) -> bytes:
            return (
                struct.pack(">I", len(data))
                + tag
                + data
                + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
            )

        png = (
            b"\x89PNG\r\n\x1a\n"
            + chunk(b"IHDR", struct.pack(">IIBBBBB", self.width, self.height, 8, 2, 0, 0, 0))
            + chunk(b"IDAT", compressed)
            + chunk(b"IEND", b"")
        )
        return BytesIO(png)
