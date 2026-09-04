import re
from pathlib import Path
from typing import Any, Dict, List, Tuple


class GroundingGuard:
    """Deterministic validation for data-grounded AI report sections and final DOCX files."""

    BLOCKED_TEXT_PATTERNS = [
        r"\[NỘI DUNG",
        r"\[CHÈN ẢNH",
        r"\[GỢI Ý",
        r"Gợi ý cho AI viết",
        r"PROMPT TỔNG HỢP",
        r"CHECKLIST TRƯỚC KHI NỘP",
        r"\bFACT_\d{3}\b",
        r"\bTODO\b",
        r"Lorem ipsum",
        r"Ignore previous instructions",
    ]
    OFF_TOPIC_TERMS = [
        "ARM",
        "x86",
        "CPU architecture",
        "server throughput",
        "network latency",
        "cloud workload",
        "GPU performance",
    ]
    CAUSAL_MARKERS = ["vì", "do ", "bởi vì", "nguyên nhân", "dẫn đến", "khiến", "because", "due to"]
    CAUSAL_SUPPORT_COLUMNS = ["kinh nghiệm", "experience", "thâm niên", "seniority", "tuổi", "age", "hiệu suất", "performance"]

    @classmethod
    def _normalize_number(cls, value: Any) -> List[float]:
        if isinstance(value, dict):
            nums: List[float] = []
            for item in value.values():
                nums.extend(cls._normalize_number(item))
            return nums
        if isinstance(value, list):
            nums = []
            for item in value:
                nums.extend(cls._normalize_number(item))
            return nums
        text = str(value)
        found = re.findall(r"-?\d+(?:[.,]\d+)?", text)
        result = []
        for raw in found:
            try:
                result.append(float(raw.replace(".", "").replace(",", ".")) if "," in raw else float(raw))
            except Exception:
                pass
        return result

    @classmethod
    def _numbers_from_text(cls, text: str) -> List[float]:
        numbers = []
        cleaned = re.sub(r"(?m)^\s*\d+(?:\.\d+){1,3}\s+", "", text or "")
        for raw in re.findall(r"(?<!FACT_)-?\d+(?:[.,]\d+)?", cleaned):
            if len(raw) == 4 and raw.startswith(("19", "20")):
                continue
            try:
                numbers.append(float(raw.replace(".", "").replace(",", ".")) if "," in raw else float(raw))
            except Exception:
                pass
        return numbers

    @classmethod
    def _fact_numbers(cls, facts: List[Dict[str, Any]]) -> List[float]:
        nums: List[float] = []
        for fact in facts:
            nums.extend(cls._normalize_number(fact.get("value")))
        return nums

    @classmethod
    def _dataset_profile_numbers(cls, profiles: List[Dict[str, Any]]) -> List[float]:
        nums: List[float] = []
        numeric_column_keys = ["sum", "min", "max", "average", "median", "count", "non_null_count", "null_count", "unique_count"]
        for profile in profiles or []:
            nums.extend(cls._fact_numbers(profile.get("verified_facts", [])))
            nums.extend(cls._normalize_number(profile.get("total_rows")))
            nums.extend(cls._normalize_number(profile.get("total_columns")))
            for sheet in profile.get("sheets", []) or []:
                nums.extend(cls._normalize_number(sheet.get("row_count")))
                nums.extend(cls._normalize_number(sheet.get("column_count")))
                for column in sheet.get("columns", []) or []:
                    for key in numeric_column_keys:
                        nums.extend(cls._normalize_number(column.get(key)))
                    for item in column.get("top_values", []) or []:
                        nums.extend(cls._normalize_number(item.get("count")))
                for grouped in sheet.get("grouped_statistics", []) or []:
                    for stat in grouped.get("statistics", []) or []:
                        nums.extend(cls._normalize_number(stat))
                for row in (sheet.get("records", []) or [])[:200]:
                    nums.extend(cls._normalize_number(row))
        return nums

    @classmethod
    def validate_section(cls, text: str, section_context: Dict[str, Any]) -> Dict[str, Any]:
        errors: List[Dict[str, Any]] = []
        warnings: List[Dict[str, Any]] = []
        text_value = text or ""
        lower = text_value.lower()

        for pattern in cls.BLOCKED_TEXT_PATTERNS:
            if re.search(pattern, text_value, flags=re.IGNORECASE):
                errors.append({"type": "TEMPLATE_LEAK", "pattern": pattern})

        for term in cls.OFF_TOPIC_TERMS:
            if term.lower() in lower:
                errors.append({"type": "OFF_TOPIC", "term": term})

        whitelist = section_context.get("allowed_entities") or {}
        allowed_entity_values = {
            str(value).lower()
            for values in whitelist.values()
            for value in values
            if value
        }
        suspicious_entities = []
        known_bad = ["IT", "Sales", "Hành chính", "Phát triển sản phẩm"]
        for entity in known_bad:
            if entity.lower() in lower and entity.lower() not in allowed_entity_values:
                suspicious_entities.append(entity)
        if suspicious_entities:
            errors.append({"type": "HALLUCINATED_ENTITY", "entities": suspicious_entities})

        fact_numbers = cls._fact_numbers(section_context.get("verified_facts", []))
        fact_numbers.extend(cls._dataset_profile_numbers(section_context.get("dataset_profiles", [])))
        generated_numbers = cls._numbers_from_text(text_value)
        unmatched = []
        for number in generated_numbers:
            if abs(number) <= 3:
                continue
            if not any(abs(number - fact) < 0.01 or (fact and abs(number - fact) / max(abs(fact), 1) < 0.001) for fact in fact_numbers):
                unmatched.append(number)
        if unmatched and fact_numbers:
            errors.append({"type": "NUMERIC_CONFLICT", "generated": unmatched[:12], "expected_values": fact_numbers[:30]})

        column_names = [str(x).lower() for x in whitelist.get("column_names", [])]
        has_causal_support = any(any(marker in col for marker in cls.CAUSAL_SUPPORT_COLUMNS) for col in column_names)
        if not has_causal_support:
            for sentence in re.split(r"(?<=[.!?。])\s+", text_value):
                if any(marker in sentence.lower() for marker in cls.CAUSAL_MARKERS):
                    errors.append({"type": "UNSUPPORTED_CAUSAL_CLAIM", "text": sentence[:220]})
                    break

        numeric_accuracy = 100 if not any(e["type"] == "NUMERIC_CONFLICT" for e in errors) else 0
        entity_accuracy = 100 if not any(e["type"] == "HALLUCINATED_ENTITY" for e in errors) else 0
        topic_relevance = 100 if not any(e["type"] == "OFF_TOPIC" for e in errors) else 0
        template_cleanliness = 100 if not any(e["type"] == "TEMPLATE_LEAK" for e in errors) else 0
        unsupported_claim_count = sum(1 for e in errors if e["type"] == "UNSUPPORTED_CAUSAL_CLAIM")
        valid = not errors
        return {
            "valid": valid,
            "errors": errors,
            "warnings": warnings,
            "scores": {
                "numeric_accuracy": numeric_accuracy,
                "entity_accuracy": entity_accuracy,
                "topic_relevance": topic_relevance,
                "template_cleanliness": template_cleanliness,
                "completeness": 100 if len(text_value.split()) >= 80 else 70,
            },
            "unsupported_claim_count": unsupported_claim_count,
            "placeholder_count": sum(1 for pattern in cls.BLOCKED_TEXT_PATTERNS if re.search(pattern, text_value, flags=re.IGNORECASE)),
        }

    @classmethod
    def final_quality_gate(cls, validations: List[Dict[str, Any]]) -> Dict[str, Any]:
        if not validations:
            return {"final": False, "reason": "NO_VALIDATION_RESULTS"}
        scores = {"numeric_accuracy": 100, "entity_accuracy": 100, "topic_relevance": 100, "template_cleanliness": 100, "completeness": 100}
        placeholder_count = 0
        unsupported_claim_count = 0
        errors: List[Dict[str, Any]] = []
        for item in validations:
            for key in scores:
                scores[key] = min(scores[key], int((item.get("scores") or {}).get(key, 100)))
            placeholder_count += int(item.get("placeholder_count") or 0)
            unsupported_claim_count += int(item.get("unsupported_claim_count") or 0)
            errors.extend(item.get("errors") or [])
        final = (
            scores["numeric_accuracy"] == 100
            and scores["entity_accuracy"] == 100
            and scores["template_cleanliness"] == 100
            and scores["topic_relevance"] >= 95
            and placeholder_count == 0
            and unsupported_claim_count == 0
        )
        return {
            "final": final,
            "scores": scores,
            "placeholder_count": placeholder_count,
            "unsupported_claim_count": unsupported_claim_count,
            "errors": errors[:60],
        }

    @classmethod
    def readiness_score(cls, validations: List[Dict[str, Any]]) -> Dict[str, Any]:
        if not validations:
            return {
                "ready": False,
                "score": 0,
                "grade": "not_checked",
                "checks": {},
                "issue_count": 1,
                "recommended_actions": ["Run quality validation before export"],
            }

        gate = cls.final_quality_gate(validations)
        checks = gate.get("scores") or {}
        base_score = int(sum(checks.values()) / max(len(checks), 1))
        errors = gate.get("errors") or []
        placeholder_count = int(gate.get("placeholder_count") or 0)
        unsupported_claim_count = int(gate.get("unsupported_claim_count") or 0)
        issue_count = len(errors) + placeholder_count + unsupported_claim_count
        score = max(0, min(100, base_score - issue_count * 12))

        if gate.get("final") and score >= 90:
            grade = "ready"
        elif score >= 75:
            grade = "review"
        elif score >= 50:
            grade = "needs_work"
        else:
            grade = "blocked"

        error_types = {str(error.get("type")) for error in errors}
        recommended_actions: List[str] = []
        if "NUMERIC_CONFLICT" in error_types:
            recommended_actions.append("Fix numeric conflicts")
        if "HALLUCINATED_ENTITY" in error_types:
            recommended_actions.append("Remove unsupported entities")
        if "OFF_TOPIC" in error_types or "FINAL_DOC_OFF_TOPIC" in error_types:
            recommended_actions.append("Remove off-topic content")
        if placeholder_count:
            recommended_actions.append("Remove placeholders and template instructions")
        if unsupported_claim_count:
            recommended_actions.append("Add evidence for causal claims")
        if not recommended_actions and not gate.get("final"):
            recommended_actions.append("Review failed quality checks")

        return {
            "ready": bool(gate.get("final")) and score >= 90,
            "score": score,
            "grade": grade,
            "checks": checks,
            "issue_count": issue_count,
            "recommended_actions": recommended_actions,
        }

    @classmethod
    def validate_docx(cls, file_path: str, topic_text: str = "") -> Dict[str, Any]:
        try:
            import docx
            document = docx.Document(file_path)
            text = "\n".join(p.text for p in document.paragraphs)
            table_empty_count = 0
            for table in document.tables:
                cells = [cell.text.strip() for row in table.rows for cell in row.cells]
                if not any(cells):
                    table_empty_count += 1
        except Exception as ex:
            return {"valid": False, "errors": [{"type": "DOCX_PARSE_FAILED", "message": str(ex)}]}

        errors = []
        for pattern in cls.BLOCKED_TEXT_PATTERNS:
            if re.search(pattern, text, flags=re.IGNORECASE):
                errors.append({"type": "FINAL_DOC_DIRTY_TEXT", "pattern": pattern})
        if "|" in text and "---" in text:
            errors.append({"type": "MARKDOWN_NOT_RENDERED"})
        fake_url_matches = re.findall(r"https?://(?:example\.com|localhost|fake|placeholder)[^\s]*", text, flags=re.IGNORECASE)
        if fake_url_matches:
            errors.append({"type": "FAKE_URL", "urls": fake_url_matches[:10]})
        if table_empty_count:
            errors.append({"type": "EMPTY_TABLE", "count": table_empty_count})
        topic_lower = (topic_text or "").lower()
        for term in cls.OFF_TOPIC_TERMS:
            term_lower = term.lower()
            if term_lower in topic_lower:
                continue
            if term_lower in text.lower():
                errors.append({"type": "FINAL_DOC_OFF_TOPIC", "term": term})
        return {
            "valid": not errors,
            "file": Path(file_path).name,
            "errors": errors,
            "word_count": len(text.split()),
        }


grounding_guard = GroundingGuard()
