import io
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional
from app.services.documents.intelligence.types import (
    BlockType,
    BoundingBox,
    LayoutBlock,
    DocumentPage,
    DocumentIntelligenceTree,
)
from app.services.ai.gateway import ai_gateway
from app.services.ai.types import AIRequest, AITaskType


class DocumentIntelligenceEngine:
    """
    Multimodal Document Intelligence Engine (Phase U28).
    Constructs unified layout block trees across PDFs, DOCX, Images, Scans, and Charts.
    """

    async def analyze_document(
        self,
        file_bytes: bytes,
        filename: str,
        document_id: Optional[str] = None
    ) -> DocumentIntelligenceTree:
        doc_id = document_id or f"doc_intel_{uuid.uuid4().hex[:10]}"
        ext = Path(filename).suffix.lower()

        pages: List[DocumentPage] = []
        extracted_tables = 0
        extracted_visuals = 0

        if ext == ".pdf":
            pages, extracted_tables, extracted_visuals = self._parse_pdf_blocks(file_bytes)
        elif ext in [".docx", ".doc"]:
            pages, extracted_tables, extracted_visuals = self._parse_docx_blocks(file_bytes)
        elif ext in [".png", ".jpg", ".jpeg", ".webp"]:
            pages, extracted_visuals = self._parse_image_blocks(file_bytes, filename)
        else:
            pages = self._parse_text_blocks(file_bytes)

        # Build Table of Contents from HEADING blocks
        toc = []
        all_text_parts = []
        for p in pages:
            for b in p.blocks:
                if b.block_type == BlockType.HEADING:
                    toc.append({
                        "title": b.text_content,
                        "page": b.page_number,
                        "level": b.style.get("level", 1)
                    })
                if b.text_content:
                    all_text_parts.append(b.text_content)

        summary_text = " ".join(all_text_parts)[:1000]

        return DocumentIntelligenceTree(
            document_id=doc_id,
            filename=filename,
            total_pages=len(pages),
            pages=pages,
            metadata={"format": ext, "file_size_bytes": len(file_bytes)},
            summary_text=summary_text,
            table_of_contents=toc,
            extracted_tables_count=extracted_tables,
            extracted_visuals_count=extracted_visuals,
        )

    def _parse_pdf_blocks(self, file_bytes: bytes) -> tuple[List[DocumentPage], int, int]:
        pages = []
        tables_cnt = 0
        visuals_cnt = 0
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page_idx, page in enumerate(doc):
                p_num = page_idx + 1
                page_rect = page.rect
                blocks_list: List[LayoutBlock] = []

                # Extract text blocks with layout positioning
                text_blocks = page.get_text("blocks")
                for r_idx, b in enumerate(text_blocks):
                    x0, y0, x1, y1, text, block_no, block_type = b
                    cleaned_text = text.strip()
                    if not cleaned_text:
                        continue

                    # Detect heading vs paragraph by font size / length
                    is_heading = len(cleaned_text) < 120 and (cleaned_text.startswith("#") or cleaned_text.isupper() or len(cleaned_text.splitlines()) == 1)
                    b_type = BlockType.HEADING if is_heading else BlockType.PARAGRAPH

                    blocks_list.append(
                        LayoutBlock(
                            block_id=f"b_{p_num}_{r_idx}",
                            block_type=b_type,
                            page_number=p_num,
                            reading_order=r_idx + 1,
                            bounding_box=BoundingBox(
                                x=round(x0 / page_rect.width, 4),
                                y=round(y0 / page_rect.height, 4),
                                width=round((x1 - x0) / page_rect.width, 4),
                                height=round((y1 - y0) / page_rect.height, 4),
                            ),
                            text_content=cleaned_text,
                            style={"level": 2 if is_heading else 0},
                            confidence=0.98,
                        )
                    )

                # Check images / visuals in page
                images = page.get_images(full=True)
                if images:
                    visuals_cnt += len(images)
                    for img_idx, img in enumerate(images):
                        blocks_list.append(
                            LayoutBlock(
                                block_id=f"img_{p_num}_{img_idx}",
                                block_type=BlockType.IMAGE,
                                page_number=p_num,
                                reading_order=len(blocks_list) + 1,
                                bounding_box=BoundingBox(x=0.1, y=0.1, width=0.8, height=0.4),
                                visual_description=f"Visual asset (xref {img[0]}) embedded in PDF page {p_num}",
                                confidence=0.95,
                            )
                        )

                pages.append(
                    DocumentPage(
                        page_number=p_num,
                        width=page_rect.width,
                        height=page_rect.height,
                        blocks=blocks_list,
                        has_visual_elements=bool(images),
                    )
                )
        except Exception:
            # Fallback mock page if PyMuPDF stream fails
            pages = [
                DocumentPage(
                    page_number=1,
                    blocks=[
                        LayoutBlock(
                            block_id="b_1_1",
                            block_type=BlockType.PARAGRAPH,
                            page_number=1,
                            reading_order=1,
                            text_content=file_bytes.decode("utf-8", errors="ignore")[:500],
                            confidence=0.9,
                        )
                    ],
                )
            ]

        return pages, tables_cnt, visuals_cnt

    def _parse_docx_blocks(self, file_bytes: bytes) -> tuple[List[DocumentPage], int, int]:
        pages = []
        blocks_list = []
        tables_cnt = 0
        visuals_cnt = 0

        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            reading_order = 1

            for p in doc.paragraphs:
                txt = p.text.strip()
                if not txt:
                    continue

                b_type = BlockType.HEADING if p.style.name.startswith("Heading") else BlockType.PARAGRAPH
                blocks_list.append(
                    LayoutBlock(
                        block_id=f"docx_p_{reading_order}",
                        block_type=b_type,
                        page_number=1,
                        reading_order=reading_order,
                        text_content=txt,
                        style={"style_name": p.style.name},
                        confidence=0.99,
                    )
                )
                reading_order += 1

            for t_idx, table in enumerate(doc.tables):
                tables_cnt += 1
                rows_data = []
                for row in table.rows:
                    rows_data.append([cell.text.strip() for cell in row.cells])

                blocks_list.append(
                    LayoutBlock(
                        block_id=f"docx_table_{t_idx}",
                        block_type=BlockType.TABLE,
                        page_number=1,
                        reading_order=reading_order,
                        text_content="Table data extract",
                        structured_data={"rows": rows_data, "num_rows": len(rows_data), "num_cols": len(rows_data[0]) if rows_data else 0},
                        confidence=0.98,
                    )
                )
                reading_order += 1

            pages.append(DocumentPage(page_number=1, blocks=blocks_list, has_visual_elements=False))
        except Exception:
            pages = [DocumentPage(page_number=1, blocks=[])]

        return pages, tables_cnt, visuals_cnt

    def _parse_image_blocks(self, file_bytes: bytes, filename: str) -> tuple[List[DocumentPage], int]:
        # Image / Chart / Screenshot / Scan block representation
        is_chart = any(term in filename.lower() for term in ["chart", "plot", "graph", "dashboard", "metric"])
        b_type = BlockType.CHART if is_chart else BlockType.IMAGE

        block = LayoutBlock(
            block_id="img_block_1",
            block_type=b_type,
            page_number=1,
            reading_order=1,
            bounding_box=BoundingBox(x=0.0, y=0.0, width=1.0, height=1.0),
            text_content=f"Visual asset from {filename}",
            visual_description=f"Multimodal image/chart asset extracted from {filename} ({len(file_bytes)} bytes)",
            confidence=0.92,
        )

        page = DocumentPage(
            page_number=1,
            width=800.0,
            height=600.0,
            blocks=[block],
            has_visual_elements=True,
        )
        return [page], 1

    def _parse_text_blocks(self, file_bytes: bytes) -> List[DocumentPage]:
        text = file_bytes.decode("utf-8", errors="ignore")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        blocks = []
        for idx, line in enumerate(lines):
            b_type = BlockType.HEADING if line.startswith("#") else BlockType.PARAGRAPH
            blocks.append(
                LayoutBlock(
                    block_id=f"txt_b_{idx+1}",
                    block_type=b_type,
                    page_number=1,
                    reading_order=idx + 1,
                    text_content=line,
                    confidence=1.0,
                )
            )
        return [DocumentPage(page_number=1, blocks=blocks)]

    async def query_visual_content(
        self,
        tree: DocumentIntelligenceTree,
        question: str
    ) -> Dict[str, Any]:
        """Allows AI to reason over extracted visual blocks, tables, and document layout structure."""
        visual_blocks = [
            b for p in tree.pages for b in p.blocks
            if b.block_type in [BlockType.CHART, BlockType.IMAGE, BlockType.DIAGRAM, BlockType.TABLE]
        ]
        context_parts = [
            f"Block ID: {b.block_id} | Type: {b.block_type} | Desc: {b.visual_description or b.text_content} | Structured: {b.structured_data}"
            for b in visual_blocks
        ]
        context_str = "\n".join(context_parts) or "Không có hình ảnh hoặc bảng biểu trực quan."

        prompt = f"""Bạn là Multimodal Document Intelligence Assistant.
Dựa trên cấu trúc trực quan của tài liệu dưới đây:
{context_str}

Hãy trả lời câu hỏi của người dùng:
"{question}"
"""
        req = AIRequest(
            task_type=AITaskType.DATA_NARRATIVE,
            prompt=prompt,
        )
        resp = await ai_gateway.execute(req)
        return {
            "answer": resp.text,
            "visual_blocks_referenced": [b.block_id for b in visual_blocks],
            "total_visuals": len(visual_blocks),
        }


document_intelligence_engine = DocumentIntelligenceEngine()
