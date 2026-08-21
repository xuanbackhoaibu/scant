import os
import re
from typing import Any, Dict, List, Optional
import docx
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocxParser:
    """Low-level Docx Parser and Style Extractor."""

    @staticmethod
    def extract_document(file_path: str) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Docx file not found: {file_path}")

        doc = docx.Document(file_path)
        paragraphs_data: List[Dict[str, Any]] = []
        headings: List[Dict[str, Any]] = []
        full_text_list: List[str] = []

        # Extract sections and margins
        sections_data: List[Dict[str, Any]] = []
        for i, section in enumerate(doc.sections):
            top_mm = round(section.top_margin.mm, 1) if section.top_margin else 20.0
            bottom_mm = round(section.bottom_margin.mm, 1) if section.bottom_margin else 20.0
            left_mm = round(section.left_margin.mm, 1) if section.left_margin else 30.0
            right_mm = round(section.right_margin.mm, 1) if section.right_margin else 20.0
            page_w_mm = round(section.page_width.mm, 1) if section.page_width else 210.0
            page_h_mm = round(section.page_height.mm, 1) if section.page_height else 297.0

            paper_type = "A4" if abs(page_w_mm - 210.0) < 6 else "Letter" if abs(page_w_mm - 215.9) < 6 else "Custom"

            sections_data.append({
                "section_index": i,
                "page_size": paper_type,
                "page_width_mm": page_w_mm,
                "page_height_mm": page_h_mm,
                "margins": {
                    "top": top_mm,
                    "bottom": bottom_mm,
                    "left": left_mm,
                    "right": right_mm,
                }
            })

        # Extract paragraphs & headings
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            full_text_list.append(text)
            style_name = p.style.name if p.style else "Normal"
            is_heading = "Heading" in style_name or "heading" in style_name.lower() or text.isupper() and len(text) < 100

            level = 1
            if "Heading 1" in style_name or "CHƯƠNG" in text.upper():
                level = 1
            elif "Heading 2" in style_name:
                level = 2
            elif "Heading 3" in style_name:
                level = 3

            p_info = {
                "text": text,
                "style": style_name,
                "is_heading": is_heading,
                "level": level if is_heading else 0
            }
            paragraphs_data.append(p_info)

            if is_heading:
                headings.append(p_info)

        # Extract tables
        tables_data: List[List[List[str]]] = []
        for table in doc.tables:
            t_rows: List[List[str]] = []
            for row in table.rows:
                t_rows.append([cell.text.strip() for cell in row.cells])
            tables_data.append(t_rows)

        full_text = "\n".join(full_text_list)

        return {
            "sections": sections_data,
            "paragraphs": paragraphs_data,
            "headings": headings,
            "tables_count": len(tables_data),
            "tables": tables_data,
            "full_text": full_text,
            "word_count": len(full_text.split()),
        }

    parse = extract_document


docx_parser = DocxParser()
