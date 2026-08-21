import docx
import pytest
import pytest_asyncio
from httpx import AsyncClient, ASGITransport
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession, async_sessionmaker
from app.main import app
from app.core.database import Base, get_db
from app.core.config import settings
from app.services.templates.template_reverse_engineering_service import template_reverse_engineer

TEST_DB_URL = "sqlite+aiosqlite:///:memory:"
test_engine = create_async_engine(TEST_DB_URL, connect_args={"check_same_thread": False})
TestAsyncSession = async_sessionmaker(bind=test_engine, class_=AsyncSession, expire_on_commit=False)


@pytest_asyncio.fixture(scope="function")
async def db_session():
    async with test_engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    async with TestAsyncSession() as session:
        yield session
    async with test_engine.begin() as conn:
        await conn.run_sync(Base.metadata.drop_all)


@pytest_asyncio.fixture(scope="function")
async def client(db_session: AsyncSession):
    async def override_get_db():
        yield db_session
    app.dependency_overrides[get_db] = override_get_db
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as c:
        yield c
    app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_reverse_engineer_docx():
    # Create sample business report docx
    doc = docx.Document()
    doc.add_heading("BÁO CÁO TÀI CHÍNH NĂM {{fiscal_year}}", level=0)
    doc.add_paragraph("Doanh nghiệp: {{company_name}}")
    doc.add_paragraph("Người lập báo cáo: {{author_name}}")
    doc.add_heading("1. Tóm Tắt Tình Hình Hoạt Động (Executive Summary)", level=1)
    doc.add_paragraph("Tổng quan doanh thu năm qua đạt mức tăng trưởng 24% so với cùng kỳ.")
    doc.add_heading("2. Báo Cáo Kết Quả Kinh Doanh", level=1)
    doc.add_paragraph("Chi tiết doanh thu và lợi nhuận gộp theo từng ngành hàng.")

    test_path = settings.TEMPLATE_DIR / "test_biz_template.docx"
    doc.save(str(test_path))

    schema = await template_reverse_engineer.reverse_engineer_docx(str(test_path))
    assert schema is not None
    assert "sections" in schema
    assert "styles" in schema
    assert "fiscal_year" in schema["explicit_placeholders"]
    assert "company_name" in schema["explicit_placeholders"]
    assert schema["styles"]["paper"] == "A4"


@pytest.mark.asyncio
async def test_reverse_engineer_api(client: AsyncClient):
    reg_res = await client.post("/api/v1/auth/register", json={
        "email": "template_admin@corp.com",
        "password": "Password123!",
        "name": "Template Admin"
    })
    token = reg_res.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}

    # Upload test docx
    test_path = settings.TEMPLATE_DIR / "test_biz_template.docx"
    with open(test_path, "rb") as f:
        files = {"file": ("test_biz_template.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        res = await client.post("/api/v1/templates/reverse-engineer", files=files, headers=headers)

    assert res.status_code == 200
    data = res.json()
    assert "styles" in data
    assert "fields" in data
