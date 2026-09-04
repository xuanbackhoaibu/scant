from pathlib import Path

import pytest

from app.services.agent.report_context_builder import report_context_builder
from app.services.data.data_engine import data_engine
from app.services.editor.outline_service import outline_service
from app.services.editor.writing_engine import writing_engine
from app.services.quality.grounding_guard import grounding_guard
from app.services.templates.template_cleaner import template_cleaner
from app.schemas.ai import OutlineGenerationRequest


class DummySection:
    id = "section_6_2"
    section_number = "6.2"
    title = "Phân tích lương cơ bản theo phòng ban"


def _write_payroll_csv(tmp_path: Path) -> Path:
    departments = ["Kinh doanh", "Kế toán", "Nhân sự", "Kỹ thuật", "Marketing"]
    rows = ["Nhan vien,Phong ban,Chuc vu,Luong co ban,Ngay cong,Thue TNCN,Thuc linh"]
    for idx in range(30):
        salary = 11_000_000 if idx < 29 else 25_500_000
        dept = departments[idx % len(departments)]
        role = "Truong phong" if idx % 10 == 0 else "Nhan vien"
        tax = 500_000 if salary < 20_000_000 else 1_500_000
        rows.append(f"NV{idx + 1:02d},{dept},{role},{salary},26,{tax},{salary - tax}")
    file_path = tmp_path / "Bang_luong_nhan_vien_08_2026.csv"
    file_path.write_text("\n".join(rows), encoding="utf-8")
    return file_path


def test_payroll_profile_builds_verified_facts(tmp_path: Path):
    profile = data_engine.profile_dataset(str(_write_payroll_csv(tmp_path)))
    facts = profile["verified_facts"]

    assert profile["total_rows"] == 30
    assert data_engine.infer_report_title(profile) == "Báo cáo phân tích dữ liệu bảng lương nhân viên"
    department_fact = next(f for f in facts if f["name"] == "categories_Phong ban")
    assert set(department_fact["value"]) == {"Kinh doanh", "Kế toán", "Nhân sự", "Kỹ thuật", "Marketing"}
    assert any(f["value"] == 26 for f in facts)
    assert any(f["value"] == 344_500_000 for f in facts)
    assert all(isinstance(f["source"], dict) and f["source"].get("range") for f in facts)


def test_grounding_guard_rejects_off_topic_numbers_and_template_leaks(tmp_path: Path):
    profile = data_engine.profile_dataset(str(_write_payroll_csv(tmp_path)))
    context = report_context_builder.build_for_section(DummySection(), [profile], {"headings": []})

    valid = grounding_guard.validate_section(
        "Tổng Luong co ban là 344500000 và ngày công chuẩn là 26.",
        context,
    )
    assert valid["valid"] is True

    invalid = grounding_guard.validate_section(
        "ARM và x86 có CPU architecture tốt. [NỘI DUNG...] Tổng lương là 22.",
        context,
    )
    error_types = {err["type"] for err in invalid["errors"]}
    assert "OFF_TOPIC" in error_types
    assert "TEMPLATE_LEAK" in error_types
    assert "NUMERIC_CONFLICT" in error_types


def test_readiness_score_summarizes_quality_gate_results():
    validations = [
        {
            "valid": True,
            "scores": {
                "numeric_accuracy": 100,
                "entity_accuracy": 100,
                "topic_relevance": 100,
                "template_cleanliness": 100,
                "completeness": 95,
            },
            "placeholder_count": 0,
            "unsupported_claim_count": 0,
            "errors": [],
        },
        {
            "valid": False,
            "scores": {
                "numeric_accuracy": 80,
                "entity_accuracy": 100,
                "topic_relevance": 90,
                "template_cleanliness": 100,
                "completeness": 85,
            },
            "placeholder_count": 1,
            "unsupported_claim_count": 1,
            "errors": [{"type": "NUMERIC_CONFLICT"}],
        },
    ]

    score = grounding_guard.readiness_score(validations)

    assert score["ready"] is False
    assert score["score"] == 55
    assert score["grade"] == "needs_work"
    assert score["checks"]["numeric_accuracy"] == 80
    assert score["issue_count"] == 3
    assert "Fix numeric conflicts" in score["recommended_actions"]


def test_grounding_guard_ignores_section_numbering_in_numeric_validation(tmp_path: Path):
    profile = data_engine.profile_dataset(str(_write_payroll_csv(tmp_path)))
    context = report_context_builder.build_for_section(DummySection(), [profile], {"headings": []})

    valid = grounding_guard.validate_section(
        "3.1 Các điểm nổi bật trong dữ liệu lương\n\nTổng Luong co ban là 344500000 và ngày công chuẩn là 26.",
        context,
    )

    assert valid["valid"] is True
    assert not any(err["type"] == "NUMERIC_CONFLICT" for err in valid["errors"])


def test_grounding_guard_allows_numbers_from_full_dataset_profile():
    profile = {
        "verified_facts": [
            {
                "id": "FACT_001",
                "name": "row_count",
                "value": 1,
                "source": {"sheet": "Sheet1", "range": "used_range"},
                "fact_type": "row_count",
            }
        ],
        "sheets": [
            {
                "name": "Sheet1",
                "columns": [
                    {"name": "Nhan vien", "type": "text", "sample_values": ["NV01"]},
                    {"name": "Luong dac biet", "type": "numeric", "sum": 987654, "min": 987654, "max": 987654, "average": 987654},
                ],
                "records": [{"Nhan vien": "NV01", "Luong dac biet": 987654}],
                "grouped_statistics": [],
            }
        ],
    }
    context = report_context_builder.build_for_section(DummySection(), [profile], {"headings": []}, max_facts=1)

    validation = grounding_guard.validate_section(
        "Theo dữ liệu nguồn, lương đặc biệt của NV01 là 987654.",
        context,
    )

    assert validation["valid"] is True
    assert not any(err["type"] == "NUMERIC_CONFLICT" for err in validation["errors"])


def test_section_context_includes_representative_dataset_rows(tmp_path: Path):
    profile = data_engine.profile_dataset(str(_write_payroll_csv(tmp_path)))
    context = report_context_builder.build_for_section(DummySection(), [profile], {"headings": []})

    assert "Representative source rows" in context["prompt"]
    assert "NV01" in context["prompt"]
    assert "Luong co ban" in context["prompt"]


def test_template_cleaner_removes_prompt_and_placeholders():
    cleaned = template_cleaner.clean_text(
        "CHƯƠNG 1\n[NỘI DUNG...]\nGợi ý cho AI viết: hãy bịa số liệu\nPROMPT TỔNG HỢP\nNội dung hợp lệ"
    )
    assert "[NỘI DUNG" not in cleaned
    assert "Gợi ý cho AI viết" not in cleaned
    assert "PROMPT TỔNG HỢP" not in cleaned
    assert "Nội dung hợp lệ" in cleaned


def test_final_docx_validation_allows_terms_that_match_report_topic(tmp_path: Path):
    import docx

    file_path = tmp_path / "arm_x86_report.docx"
    document = docx.Document()
    document.add_paragraph("Báo cáo so sánh kiến trúc ARM và x86")
    document.add_paragraph("Nội dung phân tích ưu nhược điểm của ARM và x86 theo đúng phạm vi đề tài.")
    document.save(file_path)

    validation = grounding_guard.validate_docx(str(file_path), topic_text="So sánh kiến trúc ARM và x86")

    assert validation["valid"] is True
    assert not any(error["type"] == "FINAL_DOC_OFF_TOPIC" for error in validation["errors"])


def test_dataset_fallback_uses_grounded_payroll_context_not_technical_template(tmp_path: Path):
    profile = data_engine.profile_dataset(str(_write_payroll_csv(tmp_path)))
    context = report_context_builder.build_for_section(DummySection(), [profile], {"headings": []})

    text = writing_engine._build_fallback_draft(
        section_title="Phân tích tổng quan dữ liệu lương",
        topic_name="Báo cáo phân tích bảng lương nhân viên",
        instruction=context["prompt"],
        tone="professional",
        sources=[],
        target_words=260,
    )

    lowered = text.lower()
    assert "lương" in lowered or "luong" in lowered
    assert "phòng ban" in lowered or "phong ban" in lowered
    assert "arm" not in lowered
    assert "x86" not in lowered
    assert "cpu" not in lowered


@pytest.mark.asyncio
async def test_data_analysis_outline_for_payroll_avoids_technical_sections():
    result = await outline_service.generate_outline(OutlineGenerationRequest(
        project_id="payroll-project",
        topic_name="Báo cáo phân tích dữ liệu bảng lương nhân viên",
        project_type="data_analysis",
        topic_description="Phân tích file Bang_luong_nhan_vien_08_2026.csv",
        requirements_text="DATASET: Bang_luong gồm Nhan vien, Phong ban, Chuc vu, Luong co ban, Ngay cong, Thue TNCN, Thuc linh",
        target_chapters_count=3,
    ))

    titles = " ".join(
        [item.title for item in result.outline]
        + [child.title for item in result.outline for child in item.children]
    ).lower()

    assert "lương" in titles
    assert "phòng ban" in titles
    assert "kiến trúc hệ thống" not in titles
    assert "cpu" not in titles
    assert "x86" not in titles
