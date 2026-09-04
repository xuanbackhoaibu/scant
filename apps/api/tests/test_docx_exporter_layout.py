from pathlib import Path
from types import SimpleNamespace

import docx

from app.services.exports.docx_exporter import DocxExporter


def _section(title: str, level: int, text: str = ""):
    return SimpleNamespace(
        title=title,
        level=level,
        plain_text=text,
        content_json=None,
        word_count=len(text.split()),
    )


def test_plain_docx_export_starts_each_chapter_on_new_page_and_keeps_wide_tables():
    sections = [
        _section(
            "CHƯƠNG 1: TỔNG QUAN",
            1,
            "Nội dung chương một đủ dài để kiểm tra đoạn văn bản.",
        ),
        _section(
            "CHƯƠNG 2: PHÂN TÍCH",
            1,
            "Nội dung chương hai.\n\n" + "\n".join(
                [
                    "| Tiêu chí | ARM | x86 | Nhận xét |",
                    "| --- | --- | --- | --- |",
                    "| Hiệu năng | Tốt | Rất tốt | Phù hợp nhiều kịch bản |",
                    "| Điện năng | Thấp | Cao hơn | ARM có lợi thế thiết bị di động |",
                ]
            ),
        ),
    ]

    out_path = DocxExporter.generate_docx(
        report_title="So sánh kiến trúc ARM và x86",
        topic_details={},
        sections=sections,
        sources=[],
        include_cover=False,
        include_toc=False,
        include_references=False,
    )

    generated = docx.Document(out_path)
    chapter_two_index = next(
        idx for idx, paragraph in enumerate(generated.paragraphs)
        if paragraph.text.strip() == "CHƯƠNG 2: PHÂN TÍCH"
    )
    previous_paragraphs = generated.paragraphs[:chapter_two_index]
    assert any(
        any("lastRenderedPageBreak" in run._r.xml or "w:br" in run._r.xml for run in paragraph.runs)
        for paragraph in previous_paragraphs
    )

    assert generated.tables, "Expected markdown table to become a real DOCX table"
    table = generated.tables[0]
    assert len(table.columns) == 4
    first_row_widths = [cell.width for cell in table.rows[0].cells]
    assert all(width is not None and width >= 900_000 for width in first_row_widths)

    Path(out_path).unlink(missing_ok=True)


def test_docx_export_merges_loose_markdown_table_lines_into_real_table():
    sections = [
        _section(
            "CHƯƠNG 1: BẢNG THỐNG KÊ",
            1,
            "\n\n".join(
                [
                    "| Hạng mục | Giá trị | Nhận xét |",
                    "| --- | --- | --- |",
                    "| CPU | 85% | Tải cao |",
                    "| RAM | 62% | Ổn định |",
                ]
            ),
        ),
    ]

    out_path = DocxExporter.generate_docx(
        report_title="Báo cáo hiệu năng",
        topic_details={},
        sections=sections,
        sources=[],
        include_cover=False,
        include_toc=False,
        include_references=False,
    )

    generated = docx.Document(out_path)
    assert len(generated.tables) == 1
    assert generated.tables[0].cell(0, 0).text == "Hạng mục"
    assert generated.tables[0].cell(1, 2).text == "Tải cao"

    Path(out_path).unlink(missing_ok=True)
