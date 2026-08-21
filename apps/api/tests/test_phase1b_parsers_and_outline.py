import os
import pytest
import pytest_asyncio
from pathlib import Path
import docx
import pymupdf as fitz
from httpx import AsyncClient, ASGITransport
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession, async_sessionmaker
from app.main import app
from app.core.database import Base, get_db
from app.services.documents.docx_parser import docx_parser
from app.services.templates.docx_template_analyzer import template_analyzer
from app.services.documents.pdf_parser import pdf_parser

TEST_DB_URL = "sqlite+aiosqlite:///:memory:"
test_engine = create_async_engine(TEST_DB_URL, connect_args={"check_same_thread": False})
TestAsyncSession = async_sessionmaker(bind=test_engine, class_=AsyncSession, expire_on_commit=False)


@pytest_asyncio.fixture(scope="function")
async def db_session():
    async with test_engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    async with TestAsyncSession() as session:
        yield session
    async with test_engine.begin() as conn:
        await conn.run_sync(Base.metadata.drop_all)


@pytest_asyncio.fixture(scope="function")
async def client(db_session: AsyncSession):
    async def override_get_db():
        yield db_session
    app.dependency_overrides[get_db] = override_get_db
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as c:
        yield c
    app.dependency_overrides.clear()


def test_docx_parser_and_analyzer(tmp_path: Path):
    doc_path = tmp_path / "Mau_Bao_Cao.docx"
    doc = docx.Document()
    sec = doc.sections[0]
    sec.page_width = docx.shared.Mm(210)
    sec.page_height = docx.shared.Mm(297)
    sec.left_margin = docx.shared.Mm(30)
    sec.right_margin = docx.shared.Mm(20)

    doc.add_heading("TRUONG DAI HOC BACH KHOA", level=0)
    doc.add_paragraph("Ho va ten: {{student_name}}")
    doc.add_paragraph("MSSV: {{student_id}}")
    doc.add_paragraph("Giang vien huong dan: TS. Nguyen Van B")
    doc.add_heading("CHUONG 1: TONG QUAN", level=1)
    doc.add_paragraph("Day la noi dung chuong tong quan.")
    doc.save(str(doc_path))

    # Test docx parser
    parsed = docx_parser.extract_document(str(doc_path))
    assert parsed["word_count"] > 0
    assert len(parsed["headings"]) >= 1

    # Test template analyzer
    analysis = template_analyzer.analyze_template(str(doc_path))
    assert "student_name" in analysis["explicit_placeholders"]
    assert "student_id" in analysis["explicit_placeholders"]
    assert analysis["styles"]["paper"] == "A4"
    assert analysis["styles"]["margins"]["left"] == 30.0


def test_pdf_parser(tmp_path: Path):
    pdf_path = tmp_path / "Yeu_Cau.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 72), "MUC TIEU DE TAI\nXay dung he thong hoan chinh.")
    page.insert_text((50, 100), "YEU CAU BAT BUOC\n1. Xac thuc JWT\n2. Quan ly san pham.")
    page.insert_text((50, 140), "TIEU CHI CHAM DIEM\nKien truc phan mem: 30%\nThuc nghiem: 70%")
    doc.save(str(pdf_path))
    doc.close()

    parsed = pdf_parser.extract_text_and_metadata(str(pdf_path))
    assert parsed["total_pages"] == 1
    assert "Xay dung he thong hoan chinh" in parsed["full_text"]

    summary = pdf_parser.extract_requirements_summary(parsed["full_text"])
    assert len(summary["detected_requirements"]) > 0 or len(summary["detected_objectives"]) > 0


@pytest.mark.asyncio
async def test_full_phase1b_outline_and_report_flow(client: AsyncClient):
    # 1. Register & Login
    reg_res = await client.post("/api/v1/auth/register", json={
        "email": "author@test.com",
        "password": "Password123!",
        "name": "Le Van C"
    })
    token = reg_res.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}

    # 2. Create Project
    proj_res = await client.post("/api/v1/projects", json={
        "name": "Báo cáo ASP.NET Core",
        "type": "academic",
        "topic_details": {"student_name": "Le Van C", "student_id": "20220002"}
    }, headers=headers)
    project_id = proj_res.json()["id"]

    # 3. Generate Outline via AI
    outline_res = await client.post("/api/v1/ai/generate-outline", json={
        "project_id": project_id,
        "topic_name": "Xây dựng Website Thương mại Điện tử ASP.NET Core MVC",
        "subject": "Phát triển ứng dụng Web",
        "target_chapters_count": 6
    }, headers=headers)
    assert outline_res.status_code == 200
    outline_data = outline_res.json()
    assert len(outline_data["outline"]) > 0
    assert len(outline_data["objectives"]) > 0

    # 4. Create Report from Outline
    report_res = await client.post("/api/v1/reports", json={
        "project_id": project_id,
        "title": "Báo cáo Đồ án ASP.NET Core MVC",
        "report_type": "academic",
        "outline": outline_data["outline"]
    }, headers=headers)
    assert report_res.status_code == 201
    report_data = report_res.json()
    assert len(report_data["sections"]) > 0
    first_section_id = report_data["sections"][0]["id"]

    # 5. Update Section (Autosave simulation)
    update_res = await client.put(f"/api/v1/reports/sections/{first_section_id}", json={
        "plain_text": "Nội dung cập nhật mới nhất cho phần mở đầu.",
        "status": "draft"
    }, headers=headers)
    assert update_res.status_code == 200
    assert update_res.json()["status"] == "draft"
