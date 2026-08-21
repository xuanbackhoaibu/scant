import io
import docx
import pytest
import pytest_asyncio
from httpx import AsyncClient, ASGITransport
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession, async_sessionmaker
from app.main import app
from app.core.database import Base, get_db
from app.services.documents.intelligence.document_intelligence_engine import document_intelligence_engine
from app.services.documents.intelligence.types import BlockType, DocumentIntelligenceTree

TEST_DB_URL = "sqlite+aiosqlite:///:memory:"
test_engine = create_async_engine(TEST_DB_URL, connect_args={"check_same_thread": False})
TestAsyncSession = async_sessionmaker(bind=test_engine, class_=AsyncSession, expire_on_commit=False)


@pytest_asyncio.fixture(scope="function")
async def client():
    async with test_engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)

    async def override_get_db():
        async with TestAsyncSession() as session:
            try:
                yield session
                await session.commit()
            except Exception:
                await session.rollback()
                raise

    app.dependency_overrides[get_db] = override_get_db
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as c:
        yield c

    app.dependency_overrides.clear()
    async with test_engine.begin() as conn:
        await conn.run_sync(Base.metadata.drop_all)


@pytest.mark.asyncio
async def test_docx_intelligence_tree():
    # Create sample DOCX in memory
    doc = docx.Document()
    doc.add_heading("Báo Cáo Tăng Trưởng Doanh Nghiệp 2026", level=1)
    doc.add_paragraph("Doanh thu quý 1 đạt 120 tỷ VNĐ, tăng trưởng 18% so với cùng kỳ.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Chỉ tiêu"
    table.rows[0].cells[1].text = "Giá trị"
    table.rows[1].cells[0].text = "Doanh thu"
    table.rows[1].cells[1].text = "120 tỷ"

    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    docx_bytes = docx_buf.getvalue()

    # Analyze DOCX
    tree = await document_intelligence_engine.analyze_document(
        file_bytes=docx_bytes,
        filename="growth_report.docx"
    )

    assert tree.total_pages >= 1
    assert len(tree.table_of_contents) >= 1
    assert tree.table_of_contents[0]["title"] == "Báo Cáo Tăng Trưởng Doanh Nghiệp 2026"
    assert tree.extracted_tables_count == 1

    # Verify blocks
    blocks = tree.pages[0].blocks
    assert any(b.block_type == BlockType.HEADING for b in blocks)
    assert any(b.block_type == BlockType.PARAGRAPH for b in blocks)
    assert any(b.block_type == BlockType.TABLE for b in blocks)


@pytest.mark.asyncio
async def test_image_and_chart_intelligence():
    img_bytes = b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"
    tree = await document_intelligence_engine.analyze_document(
        file_bytes=img_bytes,
        filename="revenue_breakdown_chart.png"
    )

    assert tree.total_pages == 1
    assert tree.extracted_visuals_count == 1
    assert tree.pages[0].has_visual_elements is True
    assert tree.pages[0].blocks[0].block_type == BlockType.CHART


@pytest.mark.asyncio
async def test_visual_query_reasoning():
    tree = DocumentIntelligenceTree(
        document_id="doc-test-1",
        filename="financial_dashboard.png",
        total_pages=1,
        pages=[],
        extracted_tables_count=0,
        extracted_visuals_count=1
    )
    res = await document_intelligence_engine.query_visual_content(
        tree=tree,
        question="Giải thích tỷ trọng phân bổ chi phí trong biểu đồ"
    )
    assert "answer" in res
    assert len(res["answer"]) > 0


@pytest.mark.asyncio
async def test_document_intelligence_api(client: AsyncClient):
    reg_res = await client.post("/api/v1/auth/register", json={
        "email": "intel_user@corp.com",
        "password": "Password123!",
        "name": "Intel User"
    })
    token = reg_res.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}

    # Upload file to analyze
    file_content = b"# Strategic Overview\nMarket analysis shows 25% growth."
    files = {"file": ("strategy.txt", file_content, "text/plain")}
    res = await client.post("/api/v1/document-intelligence/analyze", files=files, headers=headers)
    assert res.status_code == 200
    data = res.json()
    assert data["filename"] == "strategy.txt"
    assert len(data["pages"]) >= 1
    assert len(data["table_of_contents"]) >= 1
